
#import PyInstaller
from tkinter import *
from docx import Document 
from docx.shared import RGBColor
from docx.opc.coreprops import CoreProperties
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from tkinter import messagebox
from pandas import read_csv
import csv
import PySimpleGUI as sg
import tkinter as tk
from tkinter.filedialog import askopenfile
from tkinter import ttk
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import webbrowser
import calendar
from babel.dates import format_date, parse_date, get_day_names, get_month_names
#import babel.numbers
#from babel import numbers



global df1
data=pd.read_csv("sample.csv")  
df1=pd.read_csv("msn2.csv",sep=";")
df1=df1[["ac","ac_sn"]]
df1.values.tolist()     
aircraft=df1[["ac","ac_sn"]].values.tolist()

def getData(pn):
    
    #pn
    
    df = data
    df = df[df['PN'] == f'{pn}']
    
    description = []
    chapter = []
    part_number=[]
    for i in list(df['PN DESCRIPTION']):
        description.append(i)
    
    for j in list(df['CHAPTER']):
        chapter.append(j)
    
    part_number = pn
    
    return part_number, description[0], chapter[0]

part_number,description,chapter=getData("G4000VSVA04")


from binascii import a2b_base64
import datetime as dt
from datetime import date
from PIL import ImageTk, Image
from PyPDF2 import Transformation
from babel import Locale
from tkcalendar import DateEntry
import PIL.Image  
from PIL import ImageTk
from datetime import datetime
from docx.enum.section import WD_HEADER_FOOTER
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Mm
from docx.shared import Inches
import numpy as np
from docx.enum.table import WD_TABLE_ALIGNMENT

master=tk.Tk()
master.title("NIS APP")
W, H = 1900, 1200
master.geometry("{}x{}".format(W, H))
background_img = Image.open("flow.jpg").resize((W, H), Image.ANTIALIAS)
background_tkimg = ImageTk.PhotoImage(background_img)
canv = tk.Canvas(master)
canv.create_image(0, 0, image=background_tkimg, anchor="nw")
canv.place(x=0, y=0, relwidth=1, relheight=1)
ico =Image.open("icon.jpg")
photo =ImageTk.PhotoImage(ico)
master.wm_iconphoto(False,photo)

style=ttk.Style(master)
style.theme_use("vista")
master.wm_attributes("-alpha",0.94)

def getEntry():
    part_number, description, chapter = getData(pn_alani.get())
    text=f"{description}"
    text_1=f"{chapter}"
    des_alani.insert(0,text)
    chp_alani.insert(0,text_1)

def getTransfer():
        global text_2
        global text_3
        part_number,description,chapter=getData(pn_alani.get())     
        print("des",description)
        text_2=f"{description}"
        text_3=f"{part_number}"
        text_4=sn_alani.get()
        print("DESCRIPTION:",description)   
        desc.insert(0,text_2) 
        desc.grid(row=a-1,column=1)
        pn_area.insert(0,text_3)
        pn_area.grid(row=a-1,column=2)
        sn_area.insert(0,text_4)
        sn_area.grid(row=a-1,column=3)
        w.insert(0,text_4)
        w.grid(row=a-1,column=3)
        
        des_alani.delete(0,END)
        sn_alani.delete(0,END)
        chp_alani.delete(0,END)

def getData(pn):
    df = data
    df = df[df['PN'] == f'{pn}']
    description = []
    chapter = []
    for i in list(df['PN DESCRIPTION']):
        description.append(i)
    
    for j in list(df['CHAPTER']):
        chapter.append(j)
    part_number = pn
    return part_number, description[0], chapter[0]

#DOC HAZIRLAMA

def callback(url):
    webbrowser.open_new_tab(url)
#a=a-1   
desc_list=[]
pn_list=[]
sn_list=[] 
ac_list=[]
z_top_list=[]
fd_list=[]
tsn_list=[]
csn_list=[]

def add(btn):
        
        btn.config(bg="red")
        if nis_tipi_opsiyon.get()==var2 or nis_tipi_opsiyon.get()==var3:
                
                desc_entry=desc.get()
                desc_list.append(desc_entry)
                pn_entry=pn_area.get()
                pn_list.append(pn_entry)
                sn_entry=w.get()
                sn_list.append(sn_entry)
                ac_msn_entry=ac_msn.get()
                ac_list.append(ac_msn_entry)
                z_bottom_entry=z.get()
                z_top_list.append(z_bottom_entry)
                z_entry=z_top.get()
                z_top_list.append(z_entry)
                tod_entry=tod.get()
                fd_list.append(tod_entry)  
                fd_entry=fd.get()
                fd_list.append(fd_entry)
                x_entry=x.get()
                tsn_list.append(x_entry)
                x_bottom_entry=x_bottom.get()
                tsn_list.append(x_bottom_entry)
                q_entry=q.get()
                csn_list.append(q_entry)
                q_bottom_entry=q_bottom.get()
                csn_list.append(q_bottom_entry)
                
                
                
                print(desc_list)
                print(pn_list)
                print(sn_list)
                print(ac_list)
                print(z_top_list)
                print(fd_list)
                print(tsn_list)
                print(csn_list)




        elif nis_tipi_opsiyon.get()==var1:
                desc_entry=desc.get()
                desc_list.append(desc_entry)
                pn_entry=pn_area.get()
                pn_list.append(pn_entry)
                sn_entry=w.get()
                sn_list.append(sn_entry)



                tod_entry=tod.get()
                fd_list.append(tod_entry)  
                fd_entry=fd.get()
                fd_list.append(fd_entry)
                
                print(desc_list)
                print(pn_list)
                print(sn_list)


                print(fd_list)
        else:
                desc_entry=desc.get()
                desc_list.append(desc_entry)
                pn_entry=pn_area.get()
                pn_list.append(pn_entry)
                sn_entry=w.get()
                sn_list.append(sn_entry)

                tod_entry=tod.get()
                fd_list.append(tod_entry)  
                fd_entry=fd.get()
                fd_list.append(fd_entry)




                


                
                #x_entry=x.get()
                #tsn_list.append(x_entry)
                x_bottom_entry=x_bottom.get()
                tsn_list.append(x_bottom_entry)
                #q_entry=q.get()
                #csn_list.append(q_entry)
                q_bottom_entry=q_bottom.get()
                csn_list.append(q_bottom_entry)
                
                print(desc_list)
                print(pn_list)
                print(sn_list)
                print(tsn_list)
                print(csn_list)  



                print(fd_list)




               
                








def convert():
        
    global nis_tipi_acilir_menu
    global nis_tipi_opsiyon
    global var1
    global var2
    global var3
    global var4
    global df1
    global sign
    global ztop
        
    
                
                
    if nis_tipi_opsiyon.get()==var4:
            #if tsn_list[1]=="UNKNOWN" and csn_list[1]=="UNKNOWN":
            #        tsn_list[1]=0 
            #        csn_list[1]=0
            
            for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
            for fox in range(len(csn_list)):
                    csn_list[fox]=int(csn_list[fox])
            for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
            for delta in range(len(csn_list)):
                    csn_list[delta]=str(csn_list[delta])
            
            ztop=fd_list[0]
            print(csn_list)
            print(tsn_list)
            
            print(fd_list)
            
            
    elif len(tsn_list)>=1:
            if tsn_list[0]=="UNKNOWN" and csn_list[0]=="UNKNOWN":
                    tsn_list[0]=0 
                    csn_list[0]=0
            
            
            for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
            for fox in range(len(csn_list)):
                    csn_list[fox]=int(csn_list[fox])
                             
                    

            
            

            if len(sn_list)==1:
                    total_time=tsn_list[0]+tsn_list[1]
                    total_csn=csn_list[0]+csn_list[1]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    result = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[1]=result
                    csn_list[1]=total_csn
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                        
                    print(tsn_list)
                    print(csn_list)
            elif   pn_list[0]!=pn_list[1] and sn_list[0]!=sn_list[1] and len(sn_list)==2:
                    total_time=tsn_list[0]+tsn_list[1]
                    total_csn=csn_list[0]+csn_list[1]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    result = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[1]=result
                    csn_list[1]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time1=tsn_list[2]+tsn_list[3]
                    total_csn1=csn_list[2]+csn_list[3]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result1 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[3]=result1
                    csn_list[3]=total_csn1
                    
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                        
                    print(tsn_list)
                    print(csn_list)
                    
                    
            elif   pn_list[0]!=pn_list[1] and sn_list[0]==sn_list[1] and len(sn_list)==2:
                    total_time=tsn_list[0]+tsn_list[1]
                    total_csn=csn_list[0]+csn_list[1]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    result = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[1]=result
                    csn_list[1]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time1=tsn_list[2]+tsn_list[3]
                    total_csn1=csn_list[2]+csn_list[3]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result1 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[3]=result1
                    csn_list[3]=total_csn1
                    
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                        
                    print(tsn_list)
                    print(csn_list)
                    
                    
            elif   pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1] and len(sn_list)==2:
                    total_time=tsn_list[0]+tsn_list[1]
                    total_csn=csn_list[0]+csn_list[1]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    result = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[1]=result
                    csn_list[1]=total_csn
                    tsn_list[2]=result
                    csn_list[2]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time1=tsn_list[2]+tsn_list[3]
                    total_csn1=csn_list[2]+csn_list[3]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result1 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[3]=result1
                    csn_list[3]=total_csn1
                    
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                        
                    print(tsn_list)
                    print(csn_list)
                           
            elif  pn_list[0]==pn_list[1] and sn_list[0] == sn_list[1] and pn_list[0]==pn_list[2] and sn_list[0] == sn_list[2] and len(sn_list)==3:
                      
                    for alfa in range(len(tsn_list)):
                            tsn_list[alfa]=float(tsn_list[alfa])
                    total_time0=tsn_list[0]+tsn_list[1]
                    total_csn0=csn_list[0]+csn_list[1]
                    hours0 = int(total_time0)
                    minutes0 = int((total_time0 - hours0) * 101)
                    while minutes0 >= 60:
                            hours0 += 1
                            minutes0 -= 60
                    result0 = "{:02d}.{:02d}".format(hours0, minutes0)
                    tsn_list[1]=result0
                    csn_list[1]=total_csn0
                    tsn_list[2]=result0
                    csn_list[2]=total_csn0
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    csn_list[3]=total_csn
                    tsn_list[4]=result1
                    csn_list[4]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time1=tsn_list[4]+tsn_list[5]
                    total_csn1=csn_list[4]+csn_list[5]
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[5]=result2
                    csn_list[5]=total_csn1
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)
            elif tsn_list[0]==0 and len(sn_list)==3 and sn_list[1]==sn_list[2] and pn_list[1]==pn_list[2]:
                    total_time=tsn_list[0]+tsn_list[1]
                    total_csn=csn_list[0]+csn_list[1]
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    result = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[1]=result
                    csn_list[1]=total_csn
                    tsn_list[2]=result
                    csn_list[2]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time1=tsn_list[2]+tsn_list[3]
                    total_csn1=csn_list[2]+csn_list[3]
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[3]=result2
                    tsn_list[4]=result2
                    csn_list[3]=total_csn1
                    csn_list[4]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result3 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result3       
                    csn_list[5]=total_csn2
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)       
                    print(tsn_list)

            elif sn_list[0]==sn_list[1] and sn_list[1]!=sn_list[2] and pn_list[0]==pn_list[1] and pn_list[1]!=pn_list[2] and len(sn_list)==3:
                    total_time0=tsn_list[0]+tsn_list[1]
                    total_csn0=csn_list[0]+csn_list[1]
                    hours0 = int(total_time0)
                    minutes0 = int((total_time0 - hours0) * 101)
                    while minutes0 >= 60:
                            hours0 += 1
                            minutes0 -= 60
                    result0 = "{:02d}.{:02d}".format(hours0, minutes0)
                    tsn_list[1]=result0
                    csn_list[1]=total_csn0
                    tsn_list[2]=result0
                    csn_list[2]=total_csn0
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    csn_list[3]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)       
                    print(tsn_list)


            elif tsn_list[0]==0 and  pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1] and pn_list[1]!=pn_list[2] and sn_list[1]!=sn_list[2] and pn_list[2]==pn_list[3] and sn_list[2]==sn_list[3] and len(sn_list)==4:
                    
                    total_time0=tsn_list[0]+tsn_list[1]
                    total_csn0=csn_list[0]+csn_list[1]
                    hours0 = int(total_time0)
                    minutes0 = int((total_time0 - hours0) * 101)
                    while minutes0 >= 60:
                            hours0 += 1
                            minutes0 -= 60
                    result0 = "{:02d}.{:02d}".format(hours0, minutes0)
                    tsn_list[1]=result0
                    csn_list[1]=total_csn0
                    tsn_list[2]=result0
                    csn_list[2]=total_csn0
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    csn_list[3]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])

                    total_time1=tsn_list[4]+tsn_list[5]
                    total_csn1=csn_list[4]+csn_list[5]
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[5]=result2
                    csn_list[5]=total_csn1
                    tsn_list[6]=result2
                    csn_list[6]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])

                    total_time3=tsn_list[6]+tsn_list[7]
                    total_csn3=csn_list[6]+csn_list[7]
                    hours3 = int(total_time3)
                    minutes3 = int((total_time3 - hours3) * 101)
                    while minutes3 >= 60:
                            hours3 += 1 
                            minutes3 -= 60
                    result3 = "{:02d}.{:02d}".format(hours3, minutes3)
                    tsn_list[7]=result3
                    csn_list[7]=total_csn3
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)   
                    
            elif sn_list[0]==sn_list[1] and pn_list[0]==pn_list[1] and sn_list[1]==sn_list[2] and pn_list[1]==pn_list[2] and sn_list[2]!=sn_list[3] and pn_list[2]!=pn_list[3]and len(sn_list)==4:
                    total_time0=tsn_list[0]+tsn_list[1]
                    total_csn0=csn_list[0]+csn_list[1]
                    hours0 = int(total_time0)
                    minutes0 = int((total_time0 - hours0) * 101)
                    while minutes0 >= 60:
                            hours0 += 1
                            minutes0 -= 60
                    result0 = "{:02d}.{:02d}".format(hours0, minutes0)
                    tsn_list[1]=result0
                    csn_list[1]=total_csn0
                    tsn_list[2]=result0
                    csn_list[2]=total_csn0
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    csn_list[3]=total_csn
                    tsn_list[4]=result1
                    csn_list[4]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time1=tsn_list[4]+tsn_list[5]
                    total_csn1=csn_list[4]+csn_list[5]
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[5]=result2
                    csn_list[5]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time4=tsn_list[6]+tsn_list[7]
                    total_csn2=csn_list[6]+csn_list[7]
                    
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[7]=result5
                    csn_list[7]=total_csn2
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)         
                    
            elif tsn_list[0]==0  and sn_list[0] == sn_list[1] and sn_list[1] != sn_list[2] and sn_list[2] != sn_list[3]:
                    total_time=tsn_list[0]+tsn_list[1]
                    total_csn=csn_list[0]+csn_list[1]
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    result = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[1]=result
                    csn_list[1]=total_csn
                    tsn_list[2]=result
                    csn_list[2]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    csn_list[3]=total_csn
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)
                    
            elif tsn_list[0]==0 and sn_list[0]!=sn_list[1] and len(sn_list)!=3:
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)
            elif tsn_list[0]==0 and sn_list[0]!=sn_list[1] and sn_list[1]!=sn_list[2] and  len(sn_list)==3:
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)
            elif tsn_list[0]==0 and sn_list[2]!=sn_list[3] and sn_list[3]==sn_list[4] and sn_list[0]==sn_list[2] and len(sn_list)==5: 
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    tsn_list[4]=result1
                    csn_list[3]=total_csn
                    csn_list[4]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result4 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result4
                    csn_list[5]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time5=tsn_list[8]+tsn_list[9]
                    total_csn5=csn_list[8]+csn_list[9]
                    
                    hours5 = int(total_time5)
                    minutes5 = int((total_time5 - hours5) * 101)
                    while minutes5 >= 60:
                            hours5 += 1
                            minutes5 -= 60
                    result6 = "{:02d}.{:02d}".format(hours5, minutes5)
                    tsn_list[9]=result6
                    csn_list[9]=total_csn5
                    
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                            
                    print(csn_list)
                    print(tsn_list)       
            elif tsn_list[0]==0  and sn_list[0]!=sn_list[1] and sn_list[2]==sn_list[3]:
                    total_time3=tsn_list[4]+tsn_list[5]
                    total_csn3=csn_list[4]+csn_list[5]
                    hours3 = int(total_time3)
                    minutes3 = int((total_time3 - hours3) * 101)
                    while minutes3 >= 60:
                            hours3 += 1
                            minutes3 -= 60
                    result3 = "{:02d}.{:02d}".format(hours3, minutes3)
                    tsn_list[5]=result3
                    csn_list[5]=total_csn3
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)     
                
            elif sn_list[0]!=sn_list[1]  and  pn_list[1]!=pn_list[2] and sn_list[1]!=sn_list[2] and sn_list[2]==sn_list[3] and len(sn_list)==4:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    csn_list[3]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result4 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result4
                    tsn_list[6]=result4
                    csn_list[5]=total_csn2
                    csn_list[6]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time4=tsn_list[6]+tsn_list[7]
                    total_csn4=csn_list[6]+csn_list[7]
                    
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[7]=result5
                    csn_list[7]=total_csn4
                    

                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)   
            elif    sn_list[2]!=sn_list[3] and sn_list[3]==sn_list[4] and len(sn_list)==5:
                    
                    total_time3=tsn_list[8]+tsn_list[9]
                    total_csn3=csn_list[8]+csn_list[9]
                    hours3 = int(total_time3)
                    minutes3 = int((total_time3 - hours3) * 101)
                    while minutes3 >= 60:
                            hours3 += 1 
                            minutes3 -= 60
                    result3 = "{:02d}.{:02d}".format(hours3, minutes3)
                    tsn_list[9]=result3
                    csn_list[9]=total_csn3
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                            
                    print(csn_list)
                    print(tsn_list)
            elif    sn_list[0]==sn_list[1] and sn_list[1]!=sn_list[2] and sn_list[2]==sn_list[3] and  len(sn_list)==4:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    csn_list[1]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    csn_list[3]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result4 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result4
                    tsn_list[6]=result4
                    csn_list[5]=total_csn2
                    csn_list[6]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                                
                    total_time4=tsn_list[6]+tsn_list[7]
                    total_csn2=csn_list[6]+csn_list[7]
                    
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[7]=result5
                    csn_list[7]=total_csn2
                    
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)         
                    

            elif sn_list[0]==sn_list[3] and len(sn_list)==4:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn3=csn_list[2]+csn_list[3]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    tsn_list[4]=result1
                    csn_list[3]=total_csn3
                    csn_list[4]=total_csn3
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result4 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result4
                    tsn_list[6]=result4
                    csn_list[5]=total_csn2
                    csn_list[6]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                                
                    total_time4=tsn_list[6]+tsn_list[7]
                    total_csn5=csn_list[6]+csn_list[7]
                    
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[7]=result5
                    csn_list[7]=total_csn5
                    
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)         
            elif tsn_list[0]==0  and sn_list[0]==sn_list[4] and len(sn_list)==5:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    tsn_list[4]=result1
                    csn_list[3]=total_csn
                    csn_list[4]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result4 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result4
                    tsn_list[6]=result4
                    csn_list[5]=total_csn2
                    csn_list[6]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time4=tsn_list[6]+tsn_list[7]
                    total_csn4=csn_list[6]+csn_list[7]
                    
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[7]=result5
                    tsn_list[8]=result5
                    csn_list[7]=total_csn4
                    csn_list[8]=total_csn4
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time5=tsn_list[8]+tsn_list[9]
                    total_csn5=csn_list[8]+csn_list[9]
                    
                    hours5 = int(total_time5)
                    minutes5 = int((total_time5 - hours5) * 101)
                    while minutes5 >= 60:
                            hours5 += 1
                            minutes5 -= 60
                    result6 = "{:02d}.{:02d}".format(hours5, minutes5)
                    tsn_list[9]=result6
                    csn_list[9]=total_csn5
                    
                    
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                            
                    print(csn_list)
                    print(tsn_list)    
            elif sn_list[0]==sn_list[1] and sn_list[1]==sn_list[2] and sn_list[2]!=sn_list[3] and sn_list[3]==sn_list[4] and len(sn_list)==5:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    tsn_list[4]=result1
                    csn_list[3]=total_csn
                    csn_list[4]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result4 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result4
                    csn_list[5]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time4=tsn_list[6]+tsn_list[7]
                    total_csn4=csn_list[6]+csn_list[7]
                    
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[7]=result5
                    tsn_list[8]=result5
                    csn_list[7]=total_csn4
                    csn_list[8]=total_csn4
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time5=tsn_list[8]+tsn_list[9]
                    total_csn5=csn_list[8]+csn_list[9]
                    
                    hours5 = int(total_time5)
                    minutes5 = int((total_time5 - hours5) * 101)
                    while minutes5 >= 60:
                            hours5 += 1
                            minutes5 -= 60
                    result6 = "{:02d}.{:02d}".format(hours5, minutes5)
                    tsn_list[9]=result6
                    csn_list[9]=total_csn5
                    
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                            
                    print(csn_list)
                    print(tsn_list)    


            elif tsn_list[0]==0  and pn_list[0]==pn_list[1] and pn_list[1]==pn_list[2] and sn_list[3]==sn_list[4] and len(sn_list)==5:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    tsn_list[4]=result1
                    csn_list[3]=total_csn
                    csn_list[4]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result4 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result4
                    csn_list[5]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time4=tsn_list[6]+tsn_list[7]
                    total_csn4=csn_list[6]+csn_list[7]
                    
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[7]=result5
                    tsn_list[8]=result5
                    csn_list[7]=total_csn4
                    csn_list[8]=total_csn4
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time5=tsn_list[8]+tsn_list[9]
                    total_csn5=csn_list[8]+csn_list[9]
                    
                    hours5 = int(total_time5)
                    minutes5 = int((total_time5 - hours5) * 101)
                    while minutes5 >= 60:
                            hours5 += 1
                            minutes5 -= 60
                    result6 = "{:02d}.{:02d}".format(hours5, minutes5)
                    tsn_list[9]=result6
                    csn_list[9]=total_csn5

                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                            
                    print(csn_list)
                    print(tsn_list)    
            elif tsn_list[0]==0  and pn_list[0]==pn_list[4] and sn_list[0]==sn_list[4] and len(sn_list)==5:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    tsn_list[4]=result1
                    csn_list[3]=total_csn
                    csn_list[4]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result4 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result4
                    tsn_list[6]=result4
                    csn_list[5]=total_csn2
                    csn_list[6]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time4=tsn_list[6]+tsn_list[7]
                    total_csn4=csn_list[6]+csn_list[7]
                    
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[7]=result5
                    tsn_list[8]=result5
                    csn_list[7]=total_csn4
                    csn_list[8]=total_csn4
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time5=tsn_list[8]+tsn_list[9]
                    total_csn5=csn_list[8]+csn_list[9]
                    
                    hours5 = int(total_time5)
                    minutes5 = int((total_time5 - hours5) * 101)
                    while minutes5 >= 60:
                            hours5 += 1
                            minutes5 -= 60
                    result6 = "{:02d}.{:02d}".format(hours5, minutes5)
                    tsn_list[9]=result6
                    csn_list[9]=total_csn5

                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                            
                    print(csn_list)
                    print(tsn_list)    

            elif tsn_list[0]==0  and pn_list[0]==pn_list[5] and sn_list[0]==sn_list[5] and len(sn_list)==6:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    tsn_list[4]=result1
                    csn_list[3]=total_csn
                    csn_list[4]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result4 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result4
                    tsn_list[6]=result4
                    csn_list[5]=total_csn2
                    csn_list[6]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time4=tsn_list[6]+tsn_list[7]
                    total_csn4=csn_list[6]+csn_list[7]
                    
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[7]=result5
                    tsn_list[8]=result5
                    csn_list[7]=total_csn4
                    csn_list[8]=total_csn4
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time5=tsn_list[8]+tsn_list[9]
                    total_csn5=csn_list[8]+csn_list[9]
                    
                    hours5 = int(total_time5)
                    minutes5 = int((total_time5 - hours5) * 101)
                    while minutes5 >= 60:
                            hours5 += 1
                            minutes5 -= 60
                    result6 = "{:02d}.{:02d}".format(hours5, minutes5)
                    tsn_list[9]=result6
                    csn_list[9]=total_csn5
                    tsn_list[10]=result6
                    csn_list[10]=total_csn5
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time6=tsn_list[10]+tsn_list[11]
                    total_csn6=csn_list[10]+csn_list[11]
                    hours6 = int(total_time6)
                    minutes6 = int((total_time6 - hours6) * 101)
                    while minutes6 >= 60:
                            hours6 += 1
                            minutes6 -= 60
                    result7 = "{:02d}.{:02d}".format(hours6, minutes6)
                    tsn_list[11]=result7
                    csn_list[11]=total_csn6
                    
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                            
                    print(csn_list)
                    print(tsn_list)    
            elif tsn_list[0]==0 and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1] and pn_list[1]==pn_list[2]  and sn_list[1]==sn_list[2] and pn_list[3]==pn_list[4] and  sn_list[3]==sn_list[4] and pn_list[4]==pn_list[5] and  sn_list[4]==sn_list[5] and len(sn_list)==6:
                    
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])

                    total_time9=tsn_list[2]+tsn_list[3]
                    total_csn9=csn_list[2]+csn_list[3]
                    
                    hours9 = int(total_time9)
                    minutes9 = int((total_time9 - hours9) * 101)
                    while minutes9 >= 60:
                            hours9 += 1
                            minutes9 -= 60
                    
                    result1 = "{:02d}.{:02d}".format(hours9, minutes9)
                    tsn_list[3]=result1
                    tsn_list[4]=result1
                    csn_list[3]=total_csn9
                    csn_list[4]=total_csn9
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result4 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result4
                    csn_list[5]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])

                    total_time4=tsn_list[6]+tsn_list[7]
                    total_csn4=csn_list[6]+csn_list[7]
                    
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[7]=result5
                    tsn_list[8]=result5
                    csn_list[7]=total_csn4
                    csn_list[8]=total_csn4
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time5=tsn_list[8]+tsn_list[9]
                    total_csn5=csn_list[8]+csn_list[9]
                    
                    hours5 = int(total_time5)
                    minutes5 = int((total_time5 - hours5) * 101)
                    while minutes5 >= 60:
                            hours5 += 1
                            minutes5 -= 60
                    result6 = "{:02d}.{:02d}".format(hours5, minutes5)
                    tsn_list[9]=result6
                    csn_list[9]=total_csn5
                    tsn_list[10]=result6
                    csn_list[10]=total_csn5
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time6=tsn_list[10]+tsn_list[11]
                    total_csn6=csn_list[10]+csn_list[11]
                    hours6 = int(total_time6)
                    minutes6 = int((total_time6 - hours6) * 101)
                    while minutes6 >= 60:
                            hours6 += 1
                            minutes6 -= 60
                    result7 = "{:02d}.{:02d}".format(hours6, minutes6)
                    tsn_list[11]=result7
                    csn_list[11]=total_csn6
                    
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                            
                    print(csn_list)
                    print(tsn_list)  
            elif tsn_list[0]==0 and sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5] and len(sn_list)==6:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    csn_list[3]=total_csn

                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])


                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result4 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result4
                    tsn_list[6]=result4
                    csn_list[5]=total_csn2
                    csn_list[6]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time4=tsn_list[6]+tsn_list[7]
                    total_csn4=csn_list[6]+csn_list[7]
                    
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[7]=result5
                    csn_list[7]=total_csn4
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time5=tsn_list[8]+tsn_list[9]
                    total_csn5=csn_list[8]+csn_list[9]
                    
                    hours5 = int(total_time5)
                    minutes5 = int((total_time5 - hours5) * 101)
                    while minutes5 >= 60:
                            hours5 += 1
                            minutes5 -= 60
                    result6 = "{:02d}.{:02d}".format(hours5, minutes5)
                    tsn_list[9]=result6
                    csn_list[9]=total_csn5
                    tsn_list[10]=result6
                    csn_list[10]=total_csn5

                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time6=tsn_list[10]+tsn_list[11]
                    total_csn6=csn_list[10]+csn_list[11]
                    hours6 = int(total_time6)
                    minutes6 = int((total_time6 - hours6) * 101)
                    while minutes6 >= 60:
                            hours6 += 1
                            minutes6 -= 60
                    result7 = "{:02d}.{:02d}".format(hours6, minutes6)
                    tsn_list[11]=result7
                    csn_list[11]=total_csn6

                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)
                   

                          
            elif tsn_list[0]==0  and sn_list[0]==sn_list[6]and len(sn_list)==7:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    tsn_list[4]=result1
                    csn_list[3]=total_csn
                    csn_list[4]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result4 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result4
                    tsn_list[6]=result4
                    csn_list[5]=total_csn2
                    csn_list[6]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time4=tsn_list[6]+tsn_list[7]
                    total_csn4=csn_list[6]+csn_list[7]
                    
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[7]=result5
                    tsn_list[8]=result5
                    csn_list[7]=total_csn4
                    csn_list[8]=total_csn4
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time5=tsn_list[8]+tsn_list[9]
                    total_csn5=csn_list[8]+csn_list[9]
                    
                    hours5 = int(total_time5)
                    minutes5 = int((total_time5 - hours5) * 101)
                    while minutes5 >= 60:
                            hours5 += 1
                            minutes5 -= 60
                    result6 = "{:02d}.{:02d}".format(hours5, minutes5)
                    tsn_list[9]=result6
                    csn_list[9]=total_csn5
                    tsn_list[10]=result6
                    csn_list[10]=total_csn5
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time6=tsn_list[10]+tsn_list[11]
                    total_csn6=csn_list[10]+csn_list[11]
                    hours6 = int(total_time6)
                    minutes6 = int((total_time6 - hours6) * 101)
                    while minutes6 >= 60:
                            hours6 += 1
                            minutes6 -= 60
                    result7 = "{:02d}.{:02d}".format(hours6, minutes6)
                    tsn_list[11]=result7
                    csn_list[11]=total_csn6
                    tsn_list[12]=result7
                    csn_list[12]=total_csn6
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time7=tsn_list[12]+tsn_list[13]
                    total_csn7=csn_list[12]+csn_list[13]
                    hours7 = int(total_time7)
                    minutes7= int((total_time7 - hours7) * 101)
                    while minutes7 >= 60:
                            hours7 += 1
                            minutes7 -= 60
                    result8 = "{:02d}.{:02d}".format(hours7, minutes7)
                    tsn_list[13]=result8
                    csn_list[13]=total_csn7
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)
            elif tsn_list[0]==0  and sn_list[0]==sn_list[7]and len(sn_list)==8:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    tsn_list[4]=result1
                    csn_list[3]=total_csn
                    csn_list[4]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result4 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result4
                    tsn_list[6]=result4
                    csn_list[5]=total_csn2
                    csn_list[6]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time4=tsn_list[6]+tsn_list[7]
                    total_csn4=csn_list[6]+csn_list[7]
                    
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[7]=result5
                    tsn_list[8]=result5
                    csn_list[7]=total_csn4
                    csn_list[8]=total_csn4
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time5=tsn_list[8]+tsn_list[9]
                    total_csn5=csn_list[8]+csn_list[9]
                    
                    hours5 = int(total_time5)
                    minutes5 = int((total_time5 - hours5) * 101)
                    while minutes5 >= 60:
                            hours5 += 1
                            minutes5 -= 60
                    result6 = "{:02d}.{:02d}".format(hours5, minutes5)
                    tsn_list[9]=result6
                    csn_list[9]=total_csn5
                    tsn_list[10]=result6
                    csn_list[10]=total_csn5
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time6=tsn_list[10]+tsn_list[11]
                    total_csn6=csn_list[10]+csn_list[11]
                    hours6 = int(total_time6)
                    minutes6 = int((total_time6 - hours6) * 101)
                    while minutes6 >= 60:
                            hours6 += 1
                            minutes6 -= 60
                    result7 = "{:02d}.{:02d}".format(hours6, minutes6)
                    tsn_list[11]=result7
                    csn_list[11]=total_csn6
                    tsn_list[12]=result7
                    csn_list[12]=total_csn6
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time7=tsn_list[12]+tsn_list[13]
                    total_csn7=csn_list[12]+csn_list[13]
                    hours7 = int(total_time7)
                    minutes7= int((total_time7 - hours7) * 101)
                    while minutes7 >= 60:
                            hours7 += 1
                            minutes7 -= 60
                    result8 = "{:02d}.{:02d}".format(hours7, minutes7)
                    tsn_list[13]=result8
                    csn_list[13]=total_csn7
                    tsn_list[14]=result8
                    csn_list[14]=total_csn7
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time8=tsn_list[14]+tsn_list[15]
                    total_csn8=csn_list[14]+csn_list[15]
                    hours8 = int(total_time8)
                    minutes8= int((total_time8 - hours8) * 101)
                    while minutes8 >= 60:
                            hours8 += 1
                            minutes8 -= 60
                    result9 = "{:02d}.{:02d}".format(hours8, minutes8)
                    tsn_list[15]=result9
                    csn_list[15]=total_csn8
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)
            elif tsn_list[0]==0  and sn_list[0] == sn_list[1] and sn_list[1]==sn_list[2] and  sn_list[2] == sn_list[3] and sn_list[3]!=sn_list[4] and  sn_list[4]==sn_list[5] and sn_list[6]==sn_list[7]  and len(sn_list)==8:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    csn_list[3]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time1=tsn_list[4]+tsn_list[5]
                    total_csn1=csn_list[4]+csn_list[5]
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[5]=result2
                    csn_list[5]=total_csn1
                    tsn_list[6]=result2
                    csn_list[6]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time2=tsn_list[6]+tsn_list[7]
                    total_csn2=csn_list[6]+csn_list[7]
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result3 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[7]=result3
                    csn_list[7]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time3=tsn_list[8]+tsn_list[9]
                    total_csn3=csn_list[8]+csn_list[9]
                    hours3 = int(total_time3)
                    minutes3 = int((total_time3 - hours3) * 101)
                    while minutes3 >= 60:
                            hours3 += 1
                            minutes3 -= 60
                    result4 = "{:02d}.{:02d}".format(hours3, minutes3)
                    tsn_list[9]=result4
                    csn_list[9]=total_csn3
                    tsn_list[10]=result4
                    csn_list[10]=total_csn3
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time4=tsn_list[10]+tsn_list[11]
                    total_csn4=csn_list[10]+csn_list[11]
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[11]=result5
                    csn_list[11]=total_csn4
                    tsn_list[12]=result5
                    csn_list[12]=total_csn4
                    
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time5=tsn_list[12]+tsn_list[13]
                    total_csn5=csn_list[12]+csn_list[13]
                    hours5 = int(total_time5)
                    minutes5 = int((total_time5 - hours5) * 101)
                    while minutes5 >= 60:
                            hours5 += 1
                            minutes5 -= 60
                    result6 = "{:02d}.{:02d}".format(hours5, minutes5)
                    tsn_list[13]=result6
                    csn_list[13]=total_csn5
                    tsn_list[14]=result6
                    csn_list[14]=total_csn5
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time6=tsn_list[14]+tsn_list[15]
                    total_csn6=csn_list[14]+csn_list[15]
                    hours6 = int(total_time6)
                    minutes6 = int((total_time6 - hours6) * 101)
                    while minutes6 >= 60:
                            hours6 += 1
                            minutes6 -= 60
                    result7 = "{:02d}.{:02d}".format(hours6, minutes6)
                    tsn_list[15]=result7
                    csn_list[15]=total_csn6
                    
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)      
            elif tsn_list[0]==0  and sn_list[0] == sn_list[1] and sn_list[2] == sn_list[3] and sn_list[4]==sn_list[5] and sn_list[6]==sn_list[7]  and len(sn_list)==8:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    csn_list[3]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time1=tsn_list[4]+tsn_list[5]
                    total_csn1=csn_list[4]+csn_list[5]
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[5]=result2
                    csn_list[5]=total_csn1
                    tsn_list[6]=result2
                    csn_list[6]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time2=tsn_list[6]+tsn_list[7]
                    total_csn2=csn_list[6]+csn_list[7]
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result3 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[7]=result3
                    csn_list[7]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time3=tsn_list[8]+tsn_list[9]
                    total_csn3=csn_list[8]+csn_list[9]
                    hours3 = int(total_time3)
                    minutes3 = int((total_time3 - hours3) * 101)
                    while minutes3 >= 60:
                            hours3 += 1
                            minutes3 -= 60
                    result4 = "{:02d}.{:02d}".format(hours3, minutes3)
                    tsn_list[9]=result4
                    csn_list[9]=total_csn3
                    tsn_list[10]=result4
                    csn_list[10]=total_csn3
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time4=tsn_list[10]+tsn_list[11]
                    total_csn4=csn_list[10]+csn_list[11]
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[11]=result5
                    csn_list[11]=total_csn4
                    
                    
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time5=tsn_list[12]+tsn_list[13]
                    total_csn5=csn_list[12]+csn_list[13]
                    hours5 = int(total_time5)
                    minutes5 = int((total_time5 - hours5) * 101)
                    while minutes5 >= 60:
                            hours5 += 1
                            minutes5 -= 60
                    result6 = "{:02d}.{:02d}".format(hours5, minutes5)
                    tsn_list[13]=result6
                    csn_list[13]=total_csn5
                    tsn_list[14]=result6
                    csn_list[14]=total_csn5
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time6=tsn_list[14]+tsn_list[15]
                    total_csn6=csn_list[14]+csn_list[15]
                    hours6 = int(total_time6)
                    minutes6 = int((total_time6 - hours6) * 101)
                    while minutes6 >= 60:
                            hours6 += 1
                            minutes6 -= 60
                    result7 = "{:02d}.{:02d}".format(hours6, minutes6)
                    tsn_list[15]=result7
                    csn_list[15]=total_csn6
                    
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)      
            elif tsn_list[0]==0  and sn_list[0]==sn_list[8]and len(sn_list)==9:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    tsn_list[4]=result1
                    csn_list[3]=total_csn
                    csn_list[4]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result4 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result4
                    tsn_list[6]=result4
                    csn_list[5]=total_csn2
                    csn_list[6]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time4=tsn_list[6]+tsn_list[7]
                    total_csn4=csn_list[6]+csn_list[7]
                    
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[7]=result5
                    tsn_list[8]=result5
                    csn_list[7]=total_csn4
                    csn_list[8]=total_csn4
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time5=tsn_list[8]+tsn_list[9]
                    total_csn5=csn_list[8]+csn_list[9]
                    
                    hours5 = int(total_time5)
                    minutes5 = int((total_time5 - hours5) * 101)
                    while minutes5 >= 60:
                            hours5 += 1
                            minutes5 -= 60
                    result6 = "{:02d}.{:02d}".format(hours5, minutes5)
                    tsn_list[9]=result6
                    csn_list[9]=total_csn5
                    tsn_list[10]=result6
                    csn_list[10]=total_csn5
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time6=tsn_list[10]+tsn_list[11]
                    total_csn6=csn_list[10]+csn_list[11]
                    hours6 = int(total_time6)
                    minutes6 = int((total_time6 - hours6) * 101)
                    while minutes6 >= 60:
                            hours6 += 1
                            minutes6 -= 60
                    result7 = "{:02d}.{:02d}".format(hours6, minutes6)
                    tsn_list[11]=result7
                    csn_list[11]=total_csn6
                    tsn_list[12]=result7
                    csn_list[12]=total_csn6
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time7=tsn_list[12]+tsn_list[13]
                    total_csn7=csn_list[12]+csn_list[13]
                    hours7 = int(total_time7)
                    minutes7= int((total_time7 - hours7) * 101)
                    while minutes7 >= 60:
                            hours7 += 1
                            minutes7 -= 60
                    result8 = "{:02d}.{:02d}".format(hours7, minutes7)
                    tsn_list[13]=result8
                    csn_list[13]=total_csn7
                    tsn_list[14]=result8
                    csn_list[14]=total_csn7
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time8=tsn_list[14]+tsn_list[15]
                    total_csn8=csn_list[14]+csn_list[15]
                    hours8 = int(total_time8)
                    minutes8= int((total_time8 - hours8) * 101)
                    while minutes8 >= 60:
                            hours8 += 1
                            minutes8 -= 60
                    result9 = "{:02d}.{:02d}".format(hours8, minutes8)
                    tsn_list[15]=result9
                    csn_list[15]=total_csn8
                    tsn_list[16]=result9
                    csn_list[16]=total_csn8
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time9=tsn_list[16]+tsn_list[17]
                    total_csn9=csn_list[16]+csn_list[17]
                    hours9 = int(total_time9)
                    minutes9= int((total_time9 - hours9) * 101)
                    while minutes9 >= 60:
                            hours9 += 1
                            minutes9 -= 60
                    result10 = "{:02d}.{:02d}".format(hours9, minutes9)
                    tsn_list[17]=result10
                    csn_list[17]=total_csn9
                    
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)
            elif tsn_list[0]==0  and sn_list[0] == sn_list[1] and sn_list[2] == sn_list[3] and sn_list[4]==sn_list[5] and sn_list[6]==sn_list[7] and sn_list[8]==sn_list[9] and sn_list[0]!=sn_list[9] and len(sn_list)==10:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    csn_list[3]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time1=tsn_list[4]+tsn_list[5]
                    total_csn1=csn_list[4]+csn_list[5]
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[5]=result2
                    csn_list[5]=total_csn1
                    tsn_list[6]=result2
                    csn_list[6]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time2=tsn_list[6]+tsn_list[7]
                    total_csn2=csn_list[6]+csn_list[7]
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result3 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[7]=result3
                    csn_list[7]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time3=tsn_list[8]+tsn_list[9]
                    total_csn3=csn_list[8]+csn_list[9]
                    hours3 = int(total_time3)
                    minutes3 = int((total_time3 - hours3) * 101)
                    while minutes3 >= 60:
                            hours3 += 1
                            minutes3 -= 60
                    result4 = "{:02d}.{:02d}".format(hours3, minutes3)
                    tsn_list[9]=result4
                    csn_list[9]=total_csn3
                    tsn_list[10]=result4
                    csn_list[10]=total_csn3
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time4=tsn_list[10]+tsn_list[11]
                    total_csn4=csn_list[10]+csn_list[11]
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[11]=result5
                    csn_list[11]=total_csn4
                    
                    
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time5=tsn_list[12]+tsn_list[13]
                    total_csn5=csn_list[12]+csn_list[13]
                    hours5 = int(total_time5)
                    minutes5 = int((total_time5 - hours5) * 101)
                    while minutes5 >= 60:
                            hours5 += 1
                            minutes5 -= 60
                    result6 = "{:02d}.{:02d}".format(hours5, minutes5)
                    tsn_list[13]=result6
                    csn_list[13]=total_csn5
                    tsn_list[14]=result6
                    csn_list[14]=total_csn5
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time6=tsn_list[14]+tsn_list[15]
                    total_csn6=csn_list[14]+csn_list[15]
                    hours6 = int(total_time6)
                    minutes6 = int((total_time6 - hours6) * 101)
                    while minutes6 >= 60:
                            hours6 += 1
                            minutes6 -= 60
                    result7 = "{:02d}.{:02d}".format(hours6, minutes6)
                    tsn_list[15]=result7
                    csn_list[15]=total_csn6
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time7=tsn_list[16]+tsn_list[17]
                    total_csn7=csn_list[16]+csn_list[17]
                    hours7 = int(total_time7)
                    minutes7 = int((total_time7 - hours7) * 101)
                    while minutes7 >= 60:
                            hours7 += 1
                            minutes7 -= 60
                    result8 = "{:02d}.{:02d}".format(hours7, minutes7)
                    tsn_list[17]=result8
                    csn_list[17]=total_csn7
                    tsn_list[18]=result8
                    csn_list[18]=total_csn7
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time8=tsn_list[18]+tsn_list[19]
                    total_csn8=csn_list[18]+csn_list[19]
                    hours8 = int(total_time8)
                    minutes8 = int((total_time8 - hours8) * 101)
                    while minutes8 >= 60:
                            hours8 += 1
                            minutes8 -= 60
                    result9 = "{:02d}.{:02d}".format(hours8, minutes8)
                    tsn_list[19]=result9
                    csn_list[19]=total_csn8
                    
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)
                    
            elif tsn_list[0]==0  and sn_list[0]==sn_list[9]and len(sn_list)==10:
                    total_time1=tsn_list[0]+tsn_list[1]
                    total_csn1=csn_list[0]+csn_list[1]
                    
                    hours1 = int(total_time1)
                    minutes1 = int((total_time1 - hours1) * 101)
                    while minutes1 >= 60:
                            hours1 += 1
                            minutes1 -= 60
                    result2 = "{:02d}.{:02d}".format(hours1, minutes1)
                    tsn_list[1]=result2
                    csn_list[1]=total_csn1
                    tsn_list[2]=result2
                    csn_list[2]=total_csn1
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    total_time=tsn_list[2]+tsn_list[3]
                    total_csn=csn_list[2]+csn_list[3]
                    
                    hours = int(total_time)
                    minutes = int((total_time - hours) * 101)
                    while minutes >= 60:
                            hours += 1
                            minutes -= 60
                    
                    result1 = "{:02d}.{:02d}".format(hours, minutes)
                    tsn_list[3]=result1
                    tsn_list[4]=result1
                    csn_list[3]=total_csn
                    csn_list[4]=total_csn
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time2=tsn_list[4]+tsn_list[5]
                    total_csn2=csn_list[4]+csn_list[5]
                    
                    hours2 = int(total_time2)
                    minutes2 = int((total_time2 - hours2) * 101)
                    while minutes2 >= 60:
                            hours2 += 1
                            minutes2 -= 60
                    result4 = "{:02d}.{:02d}".format(hours2, minutes2)
                    tsn_list[5]=result4
                    tsn_list[6]=result4
                    csn_list[5]=total_csn2
                    csn_list[6]=total_csn2
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                    
                    total_time4=tsn_list[6]+tsn_list[7]
                    total_csn4=csn_list[6]+csn_list[7]
                    
                    hours4 = int(total_time4)
                    minutes4 = int((total_time4 - hours4) * 101)
                    while minutes4 >= 60:
                            hours4 += 1
                            minutes4 -= 60
                    result5 = "{:02d}.{:02d}".format(hours4, minutes4)
                    tsn_list[7]=result5
                    tsn_list[8]=result5
                    csn_list[7]=total_csn4
                    csn_list[8]=total_csn4
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time5=tsn_list[8]+tsn_list[9]
                    total_csn5=csn_list[8]+csn_list[9]
                    
                    hours5 = int(total_time5)
                    minutes5 = int((total_time5 - hours5) * 101)
                    while minutes5 >= 60:
                            hours5 += 1
                            minutes5 -= 60
                    result6 = "{:02d}.{:02d}".format(hours5, minutes5)
                    tsn_list[9]=result6
                    csn_list[9]=total_csn5
                    tsn_list[10]=result6
                    csn_list[10]=total_csn5
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time6=tsn_list[10]+tsn_list[11]
                    total_csn6=csn_list[10]+csn_list[11]
                    hours6 = int(total_time6)
                    minutes6 = int((total_time6 - hours6) * 101)
                    while minutes6 >= 60:
                            hours6 += 1
                            minutes6 -= 60
                    result7 = "{:02d}.{:02d}".format(hours6, minutes6)
                    tsn_list[11]=result7
                    csn_list[11]=total_csn6
                    tsn_list[12]=result7
                    csn_list[12]=total_csn6
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time7=tsn_list[12]+tsn_list[13]
                    total_csn7=csn_list[12]+csn_list[13]
                    hours7 = int(total_time7)
                    minutes7= int((total_time7 - hours7) * 101)
                    while minutes7 >= 60:
                            hours7 += 1
                            minutes7 -= 60
                    result8 = "{:02d}.{:02d}".format(hours7, minutes7)
                    tsn_list[13]=result8
                    csn_list[13]=total_csn7
                    tsn_list[14]=result8
                    csn_list[14]=total_csn7
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time8=tsn_list[14]+tsn_list[15]
                    total_csn8=csn_list[14]+csn_list[15]
                    hours8 = int(total_time8)
                    minutes8= int((total_time8 - hours8) * 101)
                    while minutes8 >= 60:
                            hours8 += 1
                            minutes8 -= 60
                    result9 = "{:02d}.{:02d}".format(hours8, minutes8)
                    tsn_list[15]=result9
                    csn_list[15]=total_csn8
                    tsn_list[16]=result9
                    csn_list[16]=total_csn8
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time9=tsn_list[16]+tsn_list[17]
                    total_csn9=csn_list[16]+csn_list[17]
                    hours9 = int(total_time9)
                    minutes9= int((total_time9 - hours9) * 101)
                    while minutes9 >= 60:
                            hours9 += 1
                            minutes9 -= 60
                    result10 = "{:02d}.{:02d}".format(hours9, minutes9)
                    tsn_list[17]=result10
                    csn_list[17]=total_csn9
                    tsn_list[18]=result10
                    csn_list[18]=total_csn9
                    
                    for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
                        
                    total_time10=tsn_list[18]+tsn_list[19]
                    total_csn10=csn_list[18]+csn_list[19]
                    hours10 = int(total_time10)
                    minutes10= int((total_time10 - hours10) * 101)
                    while minutes10 >= 60:
                            hours10 += 1
                            minutes10 -= 60
                    result11 = "{:02d}.{:02d}".format(hours10, minutes10)
                    tsn_list[19]=result11
                    csn_list[19]=total_csn10
                    
                    for beta in range(len(tsn_list)):
                            tsn_list[beta]=str(tsn_list[beta])
                            tsn_list[beta]=tsn_list[beta].replace(".",":")
                    for delta in range(len(csn_list)):
                            csn_list[delta]=str(csn_list[delta])
                    print(csn_list)
                    print(tsn_list)
                    
            elif tsn_list[0]==0  and sn_list[0]==sn_list[10]and len(sn_list)==11:
               total_time1=tsn_list[0]+tsn_list[1]
               total_csn1=csn_list[0]+csn_list[1]
               
               hours1 = int(total_time1)
               minutes1 = int((total_time1 - hours1) * 101)
               while minutes1 >= 60:
                       hours1 += 1
                       minutes1 -= 60
               result2 = "{:02d}.{:02d}".format(hours1, minutes1)
               tsn_list[1]=result2
               csn_list[1]=total_csn1
               tsn_list[2]=result2
               csn_list[2]=total_csn1
               for alfa in range(len(tsn_list)):
                        tsn_list[alfa]=float(tsn_list[alfa])
               
               total_time=tsn_list[2]+tsn_list[3]
               total_csn=csn_list[2]+csn_list[3]
               
               hours = int(total_time)
               minutes = int((total_time - hours) * 101)
               while minutes >= 60:
                       hours += 1
                       minutes -= 60
               
               result1 = "{:02d}.{:02d}".format(hours, minutes)
               tsn_list[3]=result1
               tsn_list[4]=result1
               csn_list[3]=total_csn
               csn_list[4]=total_csn
               for alfa in range(len(tsn_list)):
                   tsn_list[alfa]=float(tsn_list[alfa])
               
               total_time2=tsn_list[4]+tsn_list[5]
               total_csn2=csn_list[4]+csn_list[5]
               
               hours2 = int(total_time2)
               minutes2 = int((total_time2 - hours2) * 101)
               while minutes2 >= 60:
                       hours2 += 1
                       minutes2 -= 60
               result4 = "{:02d}.{:02d}".format(hours2, minutes2)
               tsn_list[5]=result4
               tsn_list[6]=result4
               csn_list[5]=total_csn2
               csn_list[6]=total_csn2
               for alfa in range(len(tsn_list)):
                   tsn_list[alfa]=float(tsn_list[alfa])
               
               total_time4=tsn_list[6]+tsn_list[7]
               total_csn4=csn_list[6]+csn_list[7]
               
               hours4 = int(total_time4)
               minutes4 = int((total_time4 - hours4) * 101)
               while minutes4 >= 60:
                       hours4 += 1
                       minutes4 -= 60
               result5 = "{:02d}.{:02d}".format(hours4, minutes4)
               tsn_list[7]=result5
               tsn_list[8]=result5
               csn_list[7]=total_csn4
               csn_list[8]=total_csn4
               for alfa in range(len(tsn_list)):
                   tsn_list[alfa]=float(tsn_list[alfa])
                   
               total_time5=tsn_list[8]+tsn_list[9]
               total_csn5=csn_list[8]+csn_list[9]
               
               hours5 = int(total_time5)
               minutes5 = int((total_time5 - hours5) * 101)
               while minutes5 >= 60:
                       hours5 += 1
                       minutes5 -= 60
               result6 = "{:02d}.{:02d}".format(hours5, minutes5)
               tsn_list[9]=result6
               csn_list[9]=total_csn5
               tsn_list[10]=result6
               csn_list[10]=total_csn5
               for alfa in range(len(tsn_list)):
                   tsn_list[alfa]=float(tsn_list[alfa])
                   
               total_time6=tsn_list[10]+tsn_list[11]
               total_csn6=csn_list[10]+csn_list[11]
               hours6 = int(total_time6)
               minutes6 = int((total_time6 - hours6) * 101)
               while minutes6 >= 60:
                       hours6 += 1
                       minutes6 -= 60
               result7 = "{:02d}.{:02d}".format(hours6, minutes6)
               tsn_list[11]=result7
               csn_list[11]=total_csn6
               tsn_list[12]=result7
               csn_list[12]=total_csn6
               for alfa in range(len(tsn_list)):
                   tsn_list[alfa]=float(tsn_list[alfa])
                   
               total_time7=tsn_list[12]+tsn_list[13]
               total_csn7=csn_list[12]+csn_list[13]
               hours7 = int(total_time7)
               minutes7= int((total_time7 - hours7) * 101)
               while minutes7 >= 60:
                       hours7 += 1
                       minutes7 -= 60
               result8 = "{:02d}.{:02d}".format(hours7, minutes7)
               tsn_list[13]=result8
               csn_list[13]=total_csn7
               tsn_list[14]=result8
               csn_list[14]=total_csn7
               for alfa in range(len(tsn_list)):
                   tsn_list[alfa]=float(tsn_list[alfa])
                   
               total_time8=tsn_list[14]+tsn_list[15]
               total_csn8=csn_list[14]+csn_list[15]
               hours8 = int(total_time8)
               minutes8= int((total_time8 - hours8) * 101)
               while minutes8 >= 60:
                       hours8 += 1
                       minutes8 -= 60
               result9 = "{:02d}.{:02d}".format(hours8, minutes8)
               tsn_list[15]=result9
               csn_list[15]=total_csn8
               tsn_list[16]=result9
               csn_list[16]=total_csn8
               for alfa in range(len(tsn_list)):
                   tsn_list[alfa]=float(tsn_list[alfa])
                   
               total_time9=tsn_list[16]+tsn_list[17]
               total_csn9=csn_list[16]+csn_list[17]
               hours9 = int(total_time9)
               minutes9= int((total_time9 - hours9) * 101)
               while minutes9 >= 60:
                       hours9 += 1
                       minutes9 -= 60
               result10 = "{:02d}.{:02d}".format(hours9, minutes9)
               tsn_list[17]=result10
               csn_list[17]=total_csn9
               tsn_list[18]=result10
               csn_list[18]=total_csn9
               
               for alfa in range(len(tsn_list)):
                   tsn_list[alfa]=float(tsn_list[alfa])
                   
               total_time10=tsn_list[18]+tsn_list[19]
               total_csn10=csn_list[18]+csn_list[19]
               hours10 = int(total_time10)
               minutes10= int((total_time10 - hours10) * 101)
               while minutes10 >= 60:
                       hours10 += 1
                       minutes10 -= 60
               result11 = "{:02d}.{:02d}".format(hours10, minutes10)
               tsn_list[19]=result11
               csn_list[19]=total_csn10
               tsn_list[20]=result11
               csn_list[20]=total_csn10
               
               for alfa in range(len(tsn_list)):
                   tsn_list[alfa]=float(tsn_list[alfa])
                   
               total_time11=tsn_list[20]+tsn_list[21]
               total_csn11=csn_list[20]+csn_list[21]
               hours11 = int(total_time11)
               minutes11= int((total_time11 - hours11) * 101)
               while minutes11 >= 60:
                       hours11 += 1
                       minutes11 -= 60
               result12 = "{:02d}.{:02d}".format(hours11, minutes11)
               tsn_list[21]=result12
               csn_list[21]=total_csn11
               
               
               for beta in range(len(tsn_list)):
                       tsn_list[beta]=str(tsn_list[beta])
                       tsn_list[beta]=tsn_list[beta].replace(".",":")
               for delta in range(len(csn_list)):
                       csn_list[delta]=str(csn_list[delta])
                    
               print(csn_list)
               print(tsn_list)
                    
            else:
                    
             print(csn_list)
             print(tsn_list)
                
                   
            ztop=fd_list[0]

            accident_list=["TC-JFN 29776", "TC-JGE 29789", "TC-JOC 1522", "TC-JGZ 35739", "TC-JSH 5546"]
            
            for i in range(len(ac_list)):
                if ac_list[i] in accident_list:
                        ac_list[i] = "INCIDENT/ACCIDENT"  
                        label_pop = Label(master, text=ac_list[i], fg="red", font=("Helvetica", 30))  
                        label_pop.pack()   
            print(ac_list)
            if len(z_top_list) == 2:
                    
                
                #z_top_list[0:2]='\033[4m' + z_top_list[0] + '\033[0m\n' +z_top_list[1] 
                z_top_list[0:2]=[' '.join(z_top_list[0:2])]
                                   
            elif len(z_top_list) == 4:
                     z_top_list[0:2]=[' '.join(z_top_list[0:2])]
                     z_top_list[1:3]=[' '.join(z_top_list[1:3])]
            elif len(z_top_list) == 6:
                        z_top_list[0:2]=[' '.join(z_top_list[0:2])]
                        z_top_list[1:3]=[' '.join(z_top_list[1:3])]
                        z_top_list[2:4]=[' '.join(z_top_list[2:4])]
                        
            elif len(z_top_list)==8:
                        z_top_list[0:2]=[' '.join(z_top_list[0:2])]
                        z_top_list[1:3]=[' '.join(z_top_list[1:3])]
                        z_top_list[2:4]=[' '.join(z_top_list[2:4])]
                        z_top_list[3:5]=[" ".join(z_top_list[3:5])]
            elif len(z_top_list)==10:
                        z_top_list[0:2]=[' '.join(z_top_list[0:2])]
                        z_top_list[1:3]=[' '.join(z_top_list[1:3])]
                        z_top_list[2:4]=[' '.join(z_top_list[2:4])]
                        z_top_list[3:5]=[" ".join(z_top_list[3:5])]
                        z_top_list[4:6]=[" ".join(z_top_list[4:6])]
            elif len(z_top_list)==12:
                        z_top_list[0:2]=[' '.join(z_top_list[0:2])]
                        z_top_list[1:3]=[' '.join(z_top_list[1:3])]
                        z_top_list[2:4]=[' '.join(z_top_list[2:4])]
                        z_top_list[3:5]=[" ".join(z_top_list[3:5])]
                        z_top_list[4:6]=[" ".join(z_top_list[4:6])]
                        z_top_list[5:7]=[" ".join(z_top_list[5:7])]
            elif len(z_top_list)==14:
                        z_top_list[0:2]=[' '.join(z_top_list[0:2])]
                        z_top_list[1:3]=[' '.join(z_top_list[1:3])]
                        z_top_list[2:4]=[' '.join(z_top_list[2:4])]
                        z_top_list[3:5]=[" ".join(z_top_list[3:5])]
                        z_top_list[4:6]=[" ".join(z_top_list[4:6])]
                        z_top_list[5:7]=[" ".join(z_top_list[5:7])]
                        z_top_list[6:8]=[" ".join(z_top_list[6:8])]
            elif len(z_top_list)==16:
                        z_top_list[0:2]=[' '.join(z_top_list[0:2])]
                        z_top_list[1:3]=[' '.join(z_top_list[1:3])]
                        z_top_list[2:4]=[' '.join(z_top_list[2:4])]
                        z_top_list[3:5]=[" ".join(z_top_list[3:5])]
                        z_top_list[4:6]=[" ".join(z_top_list[4:6])]
                        z_top_list[5:7]=[" ".join(z_top_list[5:7])]
                        z_top_list[6:8]=[" ".join(z_top_list[6:8])]
                        z_top_list[7:9]=[" ".join(z_top_list[7:9])]
            elif len(z_top_list)==18:
                        z_top_list[0:2]=[' '.join(z_top_list[0:2])]
                        z_top_list[1:3]=[' '.join(z_top_list[1:3])]
                        z_top_list[2:4]=[' '.join(z_top_list[2:4])]
                        z_top_list[3:5]=[" ".join(z_top_list[3:5])]
                        z_top_list[4:6]=[" ".join(z_top_list[4:6])]
                        z_top_list[5:7]=[" ".join(z_top_list[5:7])]
                        z_top_list[6:8]=[" ".join(z_top_list[6:8])]
                        z_top_list[7:9]=[" ".join(z_top_list[7:9])]
                        z_top_list[8:10]=[" ".join(z_top_list[8:10])]
            elif len(z_top_list)==20:
                        z_top_list[0:2]=[' '.join(z_top_list[0:2])]
                        z_top_list[1:3]=[' '.join(z_top_list[1:3])]
                        z_top_list[2:4]=[' '.join(z_top_list[2:4])]
                        z_top_list[3:5]=[" ".join(z_top_list[3:5])]
                        z_top_list[4:6]=[" ".join(z_top_list[4:6])]
                        z_top_list[5:7]=[" ".join(z_top_list[5:7])]
                        z_top_list[6:8]=[" ".join(z_top_list[6:8])]
                        z_top_list[7:9]=[" ".join(z_top_list[7:9])]
                        z_top_list[8:10]=[" ".join(z_top_list[8:10])]
                        z_top_list[9:11]=[" ".join(z_top_list[9:11])]
                        
            elif len(z_top_list)==22:
                       z_top_list[0:2]=[' '.join(z_top_list[0:2])]
                       z_top_list[1:3]=[' '.join(z_top_list[1:3])]
                       z_top_list[2:4]=[' '.join(z_top_list[2:4])]
                       z_top_list[3:5]=[" ".join(z_top_list[3:5])]
                       z_top_list[4:6]=[" ".join(z_top_list[4:6])]
                       z_top_list[5:7]=[" ".join(z_top_list[5:7])]
                       z_top_list[6:8]=[" ".join(z_top_list[6:8])]
                       z_top_list[7:9]=[" ".join(z_top_list[7:9])]
                       z_top_list[8:10]=[" ".join(z_top_list[8:10])]
                       z_top_list[9:11]=[" ".join(z_top_list[9:11])]
                       z_top_list[10:12]=[" ".join(z_top_list[10:12])]
                       
            elif len(z_top_list)==24:
                       z_top_list[0:2]=[' '.join(z_top_list[0:2])]
                       z_top_list[1:3]=[' '.join(z_top_list[1:3])]
                       z_top_list[2:4]=[' '.join(z_top_list[2:4])]
                       z_top_list[3:5]=[" ".join(z_top_list[3:5])]
                       z_top_list[4:6]=[" ".join(z_top_list[4:6])]
                       z_top_list[5:7]=[" ".join(z_top_list[5:7])]
                       z_top_list[6:8]=[" ".join(z_top_list[6:8])]
                       z_top_list[7:9]=[" ".join(z_top_list[7:9])]
                       z_top_list[8:10]=[" ".join(z_top_list[8:10])]
                       z_top_list[9:11]=[" ".join(z_top_list[9:11])]
                       z_top_list[10:12]=[" ".join(z_top_list[10:12])]
                       z_top_list[11:13]=[" ".join(z_top_list[11:13])]
                    
                
            if len(fd_list) == 2:
                    
                     
                     
                     fd_list[0:2]=[' '.join(fd_list[0:2])]
                     
            elif len(fd_list) == 4:
                    fd_list[0:2]=[' '.join(fd_list[0:2])]
                    fd_list[1:3]=[' '.join(fd_list[1:3])]
            elif len(fd_list) == 6:
                    fd_list[0:2]=[' '.join(fd_list[0:2])]
                    fd_list[1:3]=[' '.join(fd_list[1:3])]
                    fd_list[2:4]=[' '.join(fd_list[2:4])]
            elif len(fd_list)==8:
                        fd_list[0:2]=[' '.join(fd_list[0:2])]
                        fd_list[1:3]=[' '.join(fd_list[1:3])]
                        fd_list[2:4]=[' '.join(fd_list[2:4])]
                        fd_list[3:5]=[" ".join(fd_list[3:5])]
            elif len(fd_list)==10:
                        fd_list[0:2]=[' '.join(fd_list[0:2])]
                        fd_list[1:3]=[' '.join(fd_list[1:3])]
                        fd_list[2:4]=[' '.join(fd_list[2:4])]
                        fd_list[3:5]=[" ".join(fd_list[3:5])]
                        fd_list[4:6]=[" ".join(fd_list[4:6])]
            elif len(fd_list)==12:
                        fd_list[0:2]=[' '.join(fd_list[0:2])]
                        fd_list[1:3]=[' '.join(fd_list[1:3])]
                        fd_list[2:4]=[' '.join(fd_list[2:4])]
                        fd_list[3:5]=[" ".join(fd_list[3:5])]
                        fd_list[4:6]=[" ".join(fd_list[4:6])]
                        fd_list[5:7]=[" ".join(fd_list[5:7])]
            elif len(fd_list)==14:
                        fd_list[0:2]=[' '.join(fd_list[0:2])]
                        fd_list[1:3]=[' '.join(fd_list[1:3])]
                        fd_list[2:4]=[' '.join(fd_list[2:4])]
                        fd_list[3:5]=[" ".join(fd_list[3:5])]
                        fd_list[4:6]=[" ".join(fd_list[4:6])]
                        fd_list[5:7]=[" ".join(fd_list[5:7])]
                        fd_list[6:8]=[" ".join(fd_list[6:8])]
            elif len(fd_list)==16:
                        fd_list[0:2]=[' '.join(fd_list[0:2])]
                        fd_list[1:3]=[' '.join(fd_list[1:3])]
                        fd_list[2:4]=[' '.join(fd_list[2:4])]
                        fd_list[3:5]=[" ".join(fd_list[3:5])]
                        fd_list[4:6]=[" ".join(fd_list[4:6])]
                        fd_list[5:7]=[" ".join(fd_list[5:7])]
                        fd_list[6:8]=[" ".join(fd_list[6:8])]
                        fd_list[7:9]=[" ".join(fd_list[7:9])]
            elif len(fd_list)==18:
                        fd_list[0:2]=[' '.join(fd_list[0:2])]
                        fd_list[1:3]=[' '.join(fd_list[1:3])]
                        fd_list[2:4]=[' '.join(fd_list[2:4])]
                        fd_list[3:5]=[" ".join(fd_list[3:5])]
                        fd_list[4:6]=[" ".join(fd_list[4:6])]
                        fd_list[5:7]=[" ".join(fd_list[5:7])]
                        fd_list[6:8]=[" ".join(fd_list[6:8])]
                        fd_list[7:9]=[" ".join(fd_list[7:9])]
                        fd_list[8:10]=[" ".join(fd_list[8:10])]
            elif len(fd_list)==20:
                        fd_list[0:2]=[' '.join(fd_list[0:2])]
                        fd_list[1:3]=[' '.join(fd_list[1:3])]
                        fd_list[2:4]=[' '.join(fd_list[2:4])]
                        fd_list[3:5]=[" ".join(fd_list[3:5])]
                        fd_list[4:6]=[" ".join(fd_list[4:6])]
                        fd_list[5:7]=[" ".join(fd_list[5:7])]
                        fd_list[6:8]=[" ".join(fd_list[6:8])]
                        fd_list[7:9]=[" ".join(fd_list[7:9])]
                        fd_list[8:10]=[" ".join(fd_list[8:10])]
                        fd_list[9:11]=[" ".join(fd_list[9:11])]
                        
            elif len(fd_list)==22:
                       fd_list[0:2]=[' '.join(fd_list[0:2])]
                       fd_list[1:3]=[' '.join(fd_list[1:3])]
                       fd_list[2:4]=[' '.join(fd_list[2:4])]
                       fd_list[3:5]=[" ".join(fd_list[3:5])]
                       fd_list[4:6]=[" ".join(fd_list[4:6])]
                       fd_list[5:7]=[" ".join(fd_list[5:7])]
                       fd_list[6:8]=[" ".join(fd_list[6:8])]
                       fd_list[7:9]=[" ".join(fd_list[7:9])]
                       fd_list[8:10]=[" ".join(fd_list[8:10])]
                       fd_list[9:11]=[" ".join(fd_list[9:11])]
                       fd_list[10:12]=[" ".join(fd_list[10:12])]
                       
            elif len(fd_list)==24:
                       fd_list[0:2]=[' '.join(fd_list[0:2])]
                       fd_list[1:3]=[' '.join(fd_list[1:3])]
                       fd_list[2:4]=[' '.join(fd_list[2:4])]
                       fd_list[3:5]=[" ".join(fd_list[3:5])]
                       fd_list[4:6]=[" ".join(fd_list[4:6])]
                       fd_list[5:7]=[" ".join(fd_list[5:7])]
                       fd_list[6:8]=[" ".join(fd_list[6:8])]
                       fd_list[7:9]=[" ".join(fd_list[7:9])]
                       fd_list[8:10]=[" ".join(fd_list[8:10])]
                       fd_list[9:11]=[" ".join(fd_list[9:11])]
                       fd_list[10:12]=[" ".join(fd_list[10:12])]
                       fd_list[11:13]=[" ".join(fd_list[11:13])]
                    
            
                    
            #for alfa in range(len(tsn_list)):
            #            tsn_list[alfa]=float(tsn_list[alfa])
            #for fox in range(len(csn_list)):
            #        csn_list[fox]=int(csn_list[fox])
            if len(tsn_list) == 2:
                    
                    
                    tsn_list[0:2]=['  \n  '.join(tsn_list[0:2])]
                    
            elif len(tsn_list) == 4:

                tsn_list[0:2]=['      \n       '.join(tsn_list[0:2])]
                tsn_list[1:3]=['      \n      '.join(tsn_list[1:3])]
            elif len(tsn_list) == 6:
                    tsn_list[0:2]=['    \n  '.join(tsn_list[0:2])]
                    tsn_list[1:3]=['    \n  '.join(tsn_list[1:3])]
                    tsn_list[2:4]=['    \n  '.join(tsn_list[2:4])]
            elif len(tsn_list)==8:
                        tsn_list[0:2]=['  \n  '.join(tsn_list[0:2])]
                        tsn_list[1:3]=['  \n  '.join(tsn_list[1:3])]
                        tsn_list[2:4]=['  \n  '.join(tsn_list[2:4])]
                        tsn_list[3:5]=["  \n  ".join(tsn_list[3:5])]
            elif len(tsn_list)==10:
                        tsn_list[0:2]=['  \n  '.join(tsn_list[0:2])]
                        tsn_list[1:3]=['  \n  '.join(tsn_list[1:3])]
                        tsn_list[2:4]=['  \n  '.join(tsn_list[2:4])]
                        tsn_list[3:5]=["  \n  ".join(tsn_list[3:5])]
                        tsn_list[4:6]=["  \n  ".join(tsn_list[4:6])]
            elif len(tsn_list)==12:
                        tsn_list[0:2]=['  \n  '.join(tsn_list[0:2])]
                        tsn_list[1:3]=['  \n  '.join(tsn_list[1:3])]
                        tsn_list[2:4]=['  \n  '.join(tsn_list[2:4])]
                        tsn_list[3:5]=["  \n  ".join(tsn_list[3:5])]
                        tsn_list[4:6]=["  \n  ".join(tsn_list[4:6])]
                        tsn_list[5:7]=["  \n  ".join(tsn_list[5:7])]
            elif len(tsn_list)==14:
                        tsn_list[0:2]=['  \n  '.join(tsn_list[0:2])]
                        tsn_list[1:3]=['  \n  '.join(tsn_list[1:3])]
                        tsn_list[2:4]=['  \n  '.join(tsn_list[2:4])]
                        tsn_list[3:5]=["  \n  ".join(tsn_list[3:5])]
                        tsn_list[4:6]=["  \n  ".join(tsn_list[4:6])]
                        tsn_list[5:7]=["  \n  ".join(tsn_list[5:7])]
                        tsn_list[6:8]=["  \n  ".join(tsn_list[6:8])]
            elif len(tsn_list)==16:
                        tsn_list[0:2]=['  \n  '.join(tsn_list[0:2])]
                        tsn_list[1:3]=['  \n  '.join(tsn_list[1:3])]
                        tsn_list[2:4]=['  \n  '.join(tsn_list[2:4])]
                        tsn_list[3:5]=["  \n  ".join(tsn_list[3:5])]
                        tsn_list[4:6]=["  \n  ".join(tsn_list[4:6])]
                        tsn_list[5:7]=["  \n  ".join(tsn_list[5:7])]
                        tsn_list[6:8]=["  \n  ".join(tsn_list[6:8])]
                        tsn_list[7:9]=["  \n  ".join(tsn_list[7:9])]
            elif len(tsn_list)==18:
                        tsn_list[0:2]=['  \n  '.join(tsn_list[0:2])]
                        tsn_list[1:3]=['  \n  '.join(tsn_list[1:3])]
                        tsn_list[2:4]=['  \n  '.join(tsn_list[2:4])]
                        tsn_list[3:5]=["  \n  ".join(tsn_list[3:5])]
                        tsn_list[4:6]=["  \n  ".join(tsn_list[4:6])]
                        tsn_list[5:7]=["  \n  ".join(tsn_list[5:7])]
                        tsn_list[6:8]=["  \n  ".join(tsn_list[6:8])]
                        tsn_list[7:9]=["  \n  ".join(tsn_list[7:9])]
                        tsn_list[8:10]=["  \n  ".join(tsn_list[8:10])]
            elif len(tsn_list)==20:
                        tsn_list[0:2]=['  \n  '.join(tsn_list[0:2])]
                        tsn_list[1:3]=['  \n  '.join(tsn_list[1:3])]
                        tsn_list[2:4]=['  \n  '.join(tsn_list[2:4])]
                        tsn_list[3:5]=["  \n  ".join(tsn_list[3:5])]
                        tsn_list[4:6]=["  \n  ".join(tsn_list[4:6])]
                        tsn_list[5:7]=["  \n  ".join(tsn_list[5:7])]
                        tsn_list[6:8]=["  \n  ".join(tsn_list[6:8])]
                        tsn_list[7:9]=["  \n  ".join(tsn_list[7:9])]
                        tsn_list[8:10]=["  \n  ".join(tsn_list[8:10])]
                        tsn_list[9:11]=["  \n  ".join(tsn_list[9:11])]
                        
            elif len(tsn_list)==22:
                        tsn_list[0:2]=['   \n  '.join(tsn_list[0:2])]
                        tsn_list[1:3]=['   \n  '.join(tsn_list[1:3])]
                        tsn_list[2:4]=['   \n'.join(tsn_list[2:4])]
                        tsn_list[3:5]=["   \n  ".join(tsn_list[3:5])]
                        tsn_list[4:6]=["   \n  ".join(tsn_list[4:6])]
                        tsn_list[5:7]=["   \n  ".join(tsn_list[5:7])]
                        tsn_list[6:8]=["   \n  ".join(tsn_list[6:8])]
                        tsn_list[7:9]=["   \n  ".join(tsn_list[7:9])]
                        tsn_list[8:10]=["   \n  ".join(tsn_list[8:10])]
                        tsn_list[9:11]=["   \n  ".join(tsn_list[9:11])]
                        tsn_list[10:12]=["   \n  ".join(tsn_list[10:12])]
                        
            elif len(tsn_list)==24:
                       tsn_list[0:2]=['   \n  '.join(tsn_list[0:2])]
                       tsn_list[1:3]=['   \n  '.join(tsn_list[1:3])]
                       tsn_list[2:4]=['   \n  '.join(tsn_list[2:4])]
                       tsn_list[3:5]=["   \n  ".join(tsn_list[3:5])]
                       tsn_list[4:6]=["   \n  ".join(tsn_list[4:6])]
                       tsn_list[5:7]=["   \n  ".join(tsn_list[5:7])]
                       tsn_list[6:8]=["   \n  ".join(tsn_list[6:8])]
                       tsn_list[7:9]=["   \n  ".join(tsn_list[7:9])]
                       tsn_list[8:10]=["    \n  ".join(tsn_list[8:10])]
                       tsn_list[9:11]=["    \n  ".join(tsn_list[9:11])]
                       tsn_list[10:12]=["   \n  ".join(tsn_list[10:12])]
                       tsn_list[11:13]=["   \n  ".join(tsn_list[11:13])]
            elif len(tsn_list)==26:
                       tsn_list[0:2]=['   \n  '.join(tsn_list[0:2])]
                       tsn_list[1:3]=['   \n  '.join(tsn_list[1:3])]
                       tsn_list[2:4]=['   \n  '.join(tsn_list[2:4])]
                       tsn_list[3:5]=["   \n  ".join(tsn_list[3:5])]
                       tsn_list[4:6]=["   \n  ".join(tsn_list[4:6])]
                       tsn_list[5:7]=["   \n  ".join(tsn_list[5:7])]
                       tsn_list[6:8]=["   \n  ".join(tsn_list[6:8])]
                       tsn_list[7:9]=["   \n  ".join(tsn_list[7:9])]
                       tsn_list[8:10]=["    \n  ".join(tsn_list[8:10])]
                       tsn_list[9:11]=["    \n  ".join(tsn_list[9:11])]
                       tsn_list[10:12]=["   \n  ".join(tsn_list[10:12])]
                       tsn_list[11:13]=["   \n  ".join(tsn_list[11:13])]
                       tsn_list[12:14]=["   \n  ".join(tsn_list[12:14])]

            elif len(tsn_list)==28:
                       tsn_list[0:2]=['   \n  '.join(tsn_list[0:2])]
                       tsn_list[1:3]=['   \n  '.join(tsn_list[1:3])]
                       tsn_list[2:4]=['   \n  '.join(tsn_list[2:4])]
                       tsn_list[3:5]=["   \n  ".join(tsn_list[3:5])]
                       tsn_list[4:6]=["   \n  ".join(tsn_list[4:6])]
                       tsn_list[5:7]=["   \n  ".join(tsn_list[5:7])]
                       tsn_list[6:8]=["   \n  ".join(tsn_list[6:8])]
                       tsn_list[7:9]=["   \n  ".join(tsn_list[7:9])]
                       tsn_list[8:10]=["    \n  ".join(tsn_list[8:10])]
                       tsn_list[9:11]=["    \n  ".join(tsn_list[9:11])]
                       tsn_list[10:12]=["   \n  ".join(tsn_list[10:12])]
                       tsn_list[11:13]=["   \n  ".join(tsn_list[11:13])]
                       tsn_list[12:14]=["   \n  ".join(tsn_list[12:14])]
                       tsn_list[13:15]=["   \n  ".join(tsn_list[13:15])]


            
                
                    
            if len(csn_list) == 2:
            
                    csn_list[0:2]=['\n'.join(csn_list[0:2])]
            elif len(csn_list) == 4:
                    
                    csn_list[0:2]=['\n'.join(csn_list[0:2])]
                    csn_list[1:3]=['\n'.join(csn_list[1:3])]
            elif len(csn_list) == 6:
                    csn_list[0:2]=['\n'.join(csn_list[0:2])]
                    csn_list[1:3]=['\n'.join(csn_list[1:3])]
                    csn_list[2:4]=['\n'.join(csn_list[2:4])]
            elif len(csn_list)==8:
                        csn_list[0:2]=['\n'.join(csn_list[0:2])]
                        csn_list[1:3]=['\n'.join(csn_list[1:3])]
                        csn_list[2:4]=['\n'.join(csn_list[2:4])]
                        csn_list[3:5]=["\n".join(csn_list[3:5])]
            elif len(csn_list)==10:
                        csn_list[0:2]=['\n'.join(csn_list[0:2])]
                        csn_list[1:3]=['\n'.join(csn_list[1:3])]
                        csn_list[2:4]=['\n'.join(csn_list[2:4])]
                        csn_list[3:5]=["\n".join(csn_list[3:5])]
                        csn_list[4:6]=["\n".join(csn_list[4:6])]
            elif len(csn_list)==12:
                        csn_list[0:2]=['\n'.join(csn_list[0:2])]
                        csn_list[1:3]=['\n'.join(csn_list[1:3])]
                        csn_list[2:4]=['\n'.join(csn_list[2:4])]
                        csn_list[3:5]=["\n".join(csn_list[3:5])]
                        csn_list[4:6]=["\n".join(csn_list[4:6])]
                        csn_list[5:7]=["\n".join(csn_list[5:7])]
            elif len(csn_list)==14:
                        csn_list[0:2]=['\n'.join(csn_list[0:2])]
                        csn_list[1:3]=['\n'.join(csn_list[1:3])]
                        csn_list[2:4]=['\n'.join(csn_list[2:4])]
                        csn_list[3:5]=["\n".join(csn_list[3:5])]
                        csn_list[4:6]=["\n".join(csn_list[4:6])]
                        csn_list[5:7]=["\n".join(csn_list[5:7])]
                        csn_list[6:8]=["\n".join(csn_list[6:8])]
            elif len(csn_list)==16:
                        csn_list[0:2]=['\n'.join(csn_list[0:2])]
                        csn_list[1:3]=['\n'.join(csn_list[1:3])]
                        csn_list[2:4]=['\n'.join(csn_list[2:4])]
                        csn_list[3:5]=["\n".join(csn_list[3:5])]
                        csn_list[4:6]=["\n".join(csn_list[4:6])]
                        csn_list[5:7]=["\n".join(csn_list[5:7])]
                        csn_list[6:8]=["\n".join(csn_list[6:8])]
                        csn_list[7:9]=["\n".join(csn_list[7:9])]
            elif len(csn_list)==18:
                        csn_list[0:2]=['\n'.join(csn_list[0:2])]
                        csn_list[1:3]=['\n'.join(csn_list[1:3])]
                        csn_list[2:4]=['\n'.join(csn_list[2:4])]
                        csn_list[3:5]=["\n".join(csn_list[3:5])]
                        csn_list[4:6]=["\n".join(csn_list[4:6])]
                        csn_list[5:7]=["\n".join(csn_list[5:7])]
                        csn_list[6:8]=["\n".join(csn_list[6:8])]
                        csn_list[7:9]=["\n".join(csn_list[7:9])]
                        csn_list[8:10]=["\n".join(csn_list[8:10])]
            elif len(csn_list)==20:
                        csn_list[0:2]=['\n'.join(csn_list[0:2])]
                        csn_list[1:3]=['\n'.join(csn_list[1:3])]
                        csn_list[2:4]=['\n'.join(csn_list[2:4])]
                        csn_list[3:5]=["\n".join(csn_list[3:5])]
                        csn_list[4:6]=["\n".join(csn_list[4:6])]
                        csn_list[5:7]=["\n".join(csn_list[5:7])]
                        csn_list[6:8]=["\n".join(csn_list[6:8])]
                        csn_list[7:9]=["\n".join(csn_list[7:9])]
                        csn_list[8:10]=["\n".join(csn_list[8:10])]
                        csn_list[9:11]=["\n".join(csn_list[9:11])]
                        
            elif len(csn_list)==22:
                        csn_list[0:2]=['\n'.join(csn_list[0:2])]
                        csn_list[1:3]=['\n'.join(csn_list[1:3])]
                        csn_list[2:4]=['\n'.join(csn_list[2:4])]
                        csn_list[3:5]=["\n".join(csn_list[3:5])]
                        csn_list[4:6]=["\n".join(csn_list[4:6])]
                        csn_list[5:7]=["\n".join(csn_list[5:7])]
                        csn_list[6:8]=["\n".join(csn_list[6:8])]
                        csn_list[7:9]=["\n".join(csn_list[7:9])]
                        csn_list[8:10]=["\n".join(csn_list[8:10])]
                        csn_list[9:11]=["\n".join(csn_list[9:11])]
                        csn_list[10:12]=["\n".join(csn_list[10:12])]
                        
            elif len(csn_list)==24:
                        csn_list[0:2]=['\n'.join(csn_list[0:2])]
                        csn_list[1:3]=['\n'.join(csn_list[1:3])]
                        csn_list[2:4]=['\n'.join(csn_list[2:4])]
                        csn_list[3:5]=["\n".join(csn_list[3:5])]
                        csn_list[4:6]=["\n".join(csn_list[4:6])]
                        csn_list[5:7]=["\n".join(csn_list[5:7])]
                        csn_list[6:8]=["\n".join(csn_list[6:8])]
                        csn_list[7:9]=["\n".join(csn_list[7:9])]
                        csn_list[8:10]=["\n".join(csn_list[8:10])]
                        csn_list[9:11]=["\n".join(csn_list[9:11])]
                        csn_list[10:12]=["\n".join(csn_list[10:12])]
                        csn_list[11:13]=["\n".join(csn_list[11:13])]

            elif len(csn_list)==26:
                        csn_list[0:2]=['\n'.join(csn_list[0:2])]
                        csn_list[1:3]=['\n'.join(csn_list[1:3])]
                        csn_list[2:4]=['\n'.join(csn_list[2:4])]
                        csn_list[3:5]=["\n".join(csn_list[3:5])]
                        csn_list[4:6]=["\n".join(csn_list[4:6])]
                        csn_list[5:7]=["\n".join(csn_list[5:7])]
                        csn_list[6:8]=["\n".join(csn_list[6:8])]
                        csn_list[7:9]=["\n".join(csn_list[7:9])]
                        csn_list[8:10]=["\n".join(csn_list[8:10])]
                        csn_list[9:11]=["\n".join(csn_list[9:11])]
                        csn_list[10:12]=["\n".join(csn_list[10:12])]
                        csn_list[11:13]=["\n".join(csn_list[11:13])]
                        csn_list[12:14]=["\n".join(csn_list[12:14])]

            elif len(csn_list)==28:
                        csn_list[0:2]=['\n'.join(csn_list[0:2])]
                        csn_list[1:3]=['\n'.join(csn_list[1:3])]
                        csn_list[2:4]=['\n'.join(csn_list[2:4])]
                        csn_list[3:5]=["\n".join(csn_list[3:5])]
                        csn_list[4:6]=["\n".join(csn_list[4:6])]
                        csn_list[5:7]=["\n".join(csn_list[5:7])]
                        csn_list[6:8]=["\n".join(csn_list[6:8])]
                        csn_list[7:9]=["\n".join(csn_list[7:9])]
                        csn_list[8:10]=["\n".join(csn_list[8:10])]
                        csn_list[9:11]=["\n".join(csn_list[9:11])]
                        csn_list[10:12]=["\n".join(csn_list[10:12])]
                        csn_list[11:13]=["\n".join(csn_list[11:13])]
                        csn_list[12:14]=["\n".join(csn_list[12:14])]
                        csn_list[13:15]=["\n".join(csn_list[13:15])]



            print(z_top_list)
            print(fd_list)
            print(tsn_list)
            print(csn_list)
    
        
    else:    
            ztop=fd_list[0]

            
            print(desc_list)
            print(pn_list)
            print(sn_list)                    

    
    if method.get()==s1:
            
                
            
                sign="                Furkan ULCAY\t\t\t\t\t\t Şule YILMAZ\n           Component Data Control\t\t\t\t    Maintenance Records\n\t\t  Chief\t\t\t\t\t\t\t  Manager"
    elif method.get()==s2:
                sign="            Bayram Erkan ÇANKAYA\t\t\t\t\t\t Şule YILMAZ\n            Aircraft Information Control\t\t\t\t    Maintenance Records\n\t\t  Supervisor\t\t\t\t\t\t\t  Manager"
            
    elif method.get()==s3:
                sign="                Özkan AYDOĞDU\t\t\t\t\t\t           Şule YILMAZ\n           Fleet Management Data Control\t\t\t\t    Maintenance Records\n\t\t   Chief\t\t\t\t\t\t\t    Manager"
                    
    elif method.get()==s4:
                sign="                Ulaş ERGIN\t\t\t\t\t\t           Fatih NAS\n                Team Leader\t\t\t\t                         Engineering Manager\n\t\t"
    
    if nis_tipi_opsiyon.get()==var2 or nis_tipi_opsiyon.get()==var3:
            
            
            df1=pd.DataFrame.from_dict(data={"DESCRIPTION":desc_list,"PN":pn_list,"SN":sn_list,"A/C_MSN":ac_list,"TRANSACTION":z_top_list,"DATE":fd_list,"TSN":tsn_list,"CSN":csn_list},orient="index").T
            pd.set_option('display.max_rows', None)
            pd.set_option('display.max_columns', None)
    elif nis_tipi_opsiyon.get()==var1:
            df1=pd.DataFrame.from_dict(data={"DESCRIPTION":desc_list,"PN":pn_list,"SN":sn_list},orient="index").T
            pd.set_option('display.max_rows', None)
            pd.set_option('display.max_columns', None)
    else:
            df1=pd.DataFrame.from_dict(data={"DESCRIPTION":desc_list,"PN":pn_list,"SN":sn_list,"TSN":tsn_list,"CSN":csn_list},orient="index").T
            pd.set_option('display.max_rows', None)
            pd.set_option('display.max_columns', None)
    
     #THY SABLONU
    if nis_tipi_opsiyon.get()==var2 and nis_operator_opsiyon.get()==var5:
            
        
        doc=Document()
        sections=doc.sections
        
        for section in sections:
            section.top_margin=Inches(0.18)
            section.bottom_margin=Inches(0.38)
            section.left_margin=Inches(0.38)
            section.right_margin=Inches(0.58)

        section=doc.sections[0]

        

        logo=doc.add_picture("turkishairlines.png", width=Inches(1.97), height=Inches(0.64))
        logo_placement=doc.paragraphs[-1]
        logo_placement.alignment=WD_ALIGN_PARAGRAPH.CENTER
        to_date=tod.get()
        #new_to_date=to_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        #,"August").replace(".09.","September").replace(".10.","October").replace(".11."
        #,"November").replace(".12.","December")
        from_date=fd.get()
        new_from_date=from_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        ztop_date=ztop.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        body_text1=f"      This is to certify that, to the best of our knowledge, the part listed below has not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or was not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of THY operations from {ztop_date} to {new_from_date}."
        body_textif=f"      This is to certify that, to the best of our knowledge, the parts listed below have not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or were not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of THY operations from {ztop_date} to {new_from_date}."
        date_ist=nis_hazirlanis_secici.get_date()
        body_text2=f"{date_ist:%d.%m.%Y}"
        to_chp=chp_alani.get()
        
        
#line sp#ace
        for _ in range(1):
            linespace_style=doc.styles["Body Text"]
            linespace=doc.add_paragraph(style=linespace_style).add_run(" ")
            linespace_style.font.size=Pt(10)
#
# certificate template

        ist=f"İSTANBUL"
        ist_style = doc.styles['Body Text']
        ist = doc.add_paragraph(style=ist_style).add_run(ist)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ist.font.size = Pt(11)
        ist.font.bold = True
        ist.font.underline =True
        ist.font.name="Arial"

        date_ist_style=doc.styles["Normal"]
        date_ist=doc.add_paragraph(style=date_ist_style).add_run(f"{body_text2}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_ist.font.size = Pt(11)
        date_ist.font.name="Arial"
    


        body_text = f"To: WHOM IT MAY CONCERN;"

        towhom_style = doc.styles['Body Text']
        towhom = doc.add_paragraph(style=towhom_style).add_run(body_text)
        towhom.font.size = Pt(11)
        towhom.font.bold = True
        towhom.font.name="Arial"

        #for _ in range(1):
                #linespace_style = doc.styles['Body Text']
                #linespace = doc.add_paragraph(style=linespace_style).add_run()
                #linespace.font.size = 10

        heading=f"NON INCIDENT/ACCIDENT STATEMENT"
        heading_style = doc.styles['Body Text']
        head = doc.add_paragraph(style=heading_style).add_run(f'{heading}')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        head.font.size = Pt(11)
        head.font.name="Arial"
        head.font.bold = True 
        #for _ in range(1):
                #linespace_style = doc.styles['Body Text']
                #linespace = doc.add_paragraph(style=linespace_style).add_run()
                #linespace.font.size = 10

        if len(sn_list) > 1:
                
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_textif}")
                body.font.size=Pt(11)
                body.font.name="Arial"
        else:
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_text1}")
                body.font.size=Pt(11)
                body.font.name="Arial"
#Table
        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10
        table=doc.add_table(df1.shape[0]+1, df1.shape[1],style="Light List Accent 1")
        table.style.font.size=Pt(9)
        table.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.border=True
        #def ChangeWidthOfTable(table,width,column):
        #        for columnVariable in range(0,column):
        #                for cell in table.columns[columnVariable].cells:
        #                        cell.width=width(15)
        
        if len(desc_list)==1 and len(pn_list)==1 and len(sn_list)==1:
                
                for k in range(df1.shape[-1]):
                        table.cell(0,k).text=df1.columns[k]
                for l in range(df1.shape[0]):
                        for k in range(df1.shape[-1]):
                                table.cell(l+1,k).text=str(df1.values[l,k])
                                
        elif len(desc_list)==2 and len(pn_list)==2 and len(sn_list)==2:
                if desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==3 and len(pn_list)==3 and len(sn_list)==3:
                if desc_list[0]==desc_list[2] and pn_list[0]==pn_list[2] and sn_list[0]==sn_list[2]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
                elif desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
                elif sn_list[1]==sn_list[2] and pn_list[1]==pn_list[2]:
                        rowa=table.rows[2]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
                
        elif len(desc_list)==4 and len(pn_list)==4 and len(sn_list)==4:
                if desc_list[0]==desc_list[3] and pn_list[0]==pn_list[3] and sn_list[0]==sn_list[3]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[4]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                elif  pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1] and pn_list[2]==pn_list[3] and sn_list[2]==sn_list[3]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
                elif desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                elif sn_list[0]==sn_list[2] and sn_list[2]!=sn_list[3]:
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])


                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==5 and len(pn_list)==5 and len(sn_list)==5:
                if desc_list[0]==desc_list[4] and pn_list[0]==pn_list[4] and sn_list[0]==sn_list[4]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[5]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])


                elif sn_list[0]==sn_list[2] and sn_list[3]==sn_list[4]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[4]
                        rowd=table.rows[5]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
        elif len(desc_list)==6 and len(pn_list)==6 and len(sn_list)==6:
                if desc_list[0]==desc_list[5] and pn_list[0]==pn_list[5] and sn_list[0]==sn_list[5]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[6]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                

                elif sn_list[0]==sn_list[2] and sn_list[3]==sn_list[5]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowc=table.rows[3]
                        #rowd=table.rows[5]
                        #rowc.cells[0].merge(rowd.cells[0])
                        #rowc.cells[1].merge(rowd.cells[1])
                        #rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[4]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])


                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])



                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==7 and len(pn_list)==7 and len(sn_list)==7:
                if desc_list[0]==desc_list[6] and pn_list[0]==pn_list[6] and sn_list[0]==sn_list[6]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[7]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])


                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==8 and len(pn_list)==8 and len(sn_list)==8:
                if desc_list[0]==desc_list[7] and pn_list[0]==pn_list[7] and sn_list[0]==sn_list[7]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[8]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[3]!=sn_list[4] and  sn_list[4]==sn_list[5] and sn_list[6]==sn_list[7]:
                        rowa=table.rows[1]
                        rowb=table.rows[4]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[5]
                        rowd=table.rows[8]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])


                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5] and sn_list[6]==sn_list[7]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])
                        rowg=table.rows[7]
                        rowh=table.rows[8]
                        rowg.cells[0].merge(rowh.cells[0])
                        rowg.cells[1].merge(rowh.cells[1])
                        rowg.cells[2].merge(rowh.cells[2])
                        

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==9 and len(pn_list)==9 and len(sn_list)==9:
                if desc_list[0]==desc_list[8] and pn_list[0]==pn_list[8] and sn_list[0]==sn_list[8]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[9]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
        elif len(desc_list)==10 and len(pn_list)==10 and len(sn_list)==10:
                if desc_list[0]==desc_list[9] and pn_list[0]==pn_list[9] and sn_list[0]==sn_list[9]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[10]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                
                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5] and sn_list[6]==sn_list[7] and sn_list[8]==sn_list[9]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])
                        rowg=table.rows[7]
                        rowh=table.rows[8]
                        rowg.cells[0].merge(rowh.cells[0])
                        rowg.cells[1].merge(rowh.cells[1])
                        rowg.cells[2].merge(rowh.cells[2])
                        rowj=table.rows[9]
                        rowk=table.rows[10]
                        rowj.cells[0].merge(rowk.cells[0])
                        rowj.cells[1].merge(rowk.cells[1])
                        rowj.cells[2].merge(rowk.cells[2])
                        
                        
                        
                        
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==11 and len(pn_list)==11 and len(sn_list)==11:
                if desc_list[0]==desc_list[10] and pn_list[0]==pn_list[10] and sn_list[0]==sn_list[10]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[11]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==12 and len(pn_list)==12 and len(sn_list)==12:
                if desc_list[0]==desc_list[11] and pn_list[0]==pn_list[11] and sn_list[0]==sn_list[11]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[12]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
        else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])                

        for _ in range(2):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10
        
        tai_text=f"TURKISH AIRLINES INC."
        tai_style=doc.styles["Body Text"]
        tai=doc.add_paragraph(style=tai_style).add_run(f"{tai_text}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tai.font.size=Pt(12)
        tai.font.name="Arial"
        tai.font.bold= True

        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        chp_style=doc.styles["Body Text"]
        chp=doc.add_paragraph(style=chp_style).add_run(f"{sign}")
        chp.font.size = Pt(11)
        chp.font.name="Arial"
        chp.font.bold= True

        for _ in range(14):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10
        
        footer_section=doc.add_picture("footer1.png", width=Inches(5.00),height=Inches(0.61))
        footer_thy=doc.paragraphs[-1]
        footer_thy.alignment= WD_ALIGN_PARAGRAPH.CENTER

    
    
        doc.save("word.docx")

    elif nis_tipi_opsiyon.get()==var1 and nis_operator_opsiyon.get()==var5:
            
        

        doc=Document()
        sections=doc.sections


        for section in sections:
            section.top_margin=Inches(0.18)
            section.bottom_margin=Inches(0.38)
            section.left_margin=Inches(0.38)
            section.right_margin=Inches(0.58)

        section=doc.sections[0]
        logo=doc.add_picture("turkishairlines.png", width=Inches(1.97), height=Inches(0.64))
        logo_placement=doc.paragraphs[-1]
        logo_placement.alignment=WD_ALIGN_PARAGRAPH.CENTER
    
        #to_date=tod.get()
        #new_to_date=to_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        #,"August").replace(".09.","September").replace(".10.","October").replace(".11."
        #,"November").replace(".12.","December")
        #from_date=fd.get()
        #new_from_date=from_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        #,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        #body_text1=f"      This is to certify that, to the best of our knowledge, the part listed below has not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or was not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of THY operations."
        #body_textif=f"      This is to certify that, to the best of our knowledge, the parts listed below have not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or were not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of THY operations."
        #date_ist=nis_hazirlanis_secici.get_date()
        #body_text2=f"{date_ist:%d.%m.%Y}"
        #to_chp=chp_alani.get()
        



        to_date=tod.get()
        #new_to_date=to_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        #,"August").replace(".09.","September").replace(".10.","October").replace(".11."
        #,"November").replace(".12.","December")
        from_date=fd.get()
        new_from_date=from_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        ztop_date=ztop.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        body_text1=f"      This is to certify that, to the best of our knowledge, the part listed below has not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or was not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of THY operations from {ztop_date} to {new_from_date}."
        body_textif=f"      This is to certify that, to the best of our knowledge, the parts listed below have not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or were not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of THY operations from {ztop_date} to {new_from_date}."
        paragraph1=f"      I hereby certify that, to the best of my knowledge, during the period stated above:\n\t1. No subject part installed has been;\n\n\t\ta. damaged during, or identified as the root cause of, a reportable incident or \t\t\taccident as defined by Annex 13 to the Chicago Convention, other than the \t\t\tfollowing:\n\n\t\t\t-Write The Occurence Report Note\n\n\t\tb.subjected to severe stress or heat such as in a major engine failure, accident,or\n\t\t fire or has been submersed in salt water,\n\n\t\t unless its airworthiness status was re-established by an approved maintenance\n\t\t organisation in accordance with the applicable airworthiness regulations and\n\t\t instructions of the type certificate holder and-or supplemental type certificate\n\t\t holder and-or OEM of the part, and supported by an authorised airworthiness\n\t\t release certificate.\n\n\t2. No part has been installed on the aircraft which was obtained from a military source or\n\t was previously fitted to a state aircraft as deemed by Article 3 of the Chicago\n\t Convention. "
        date_ist=nis_hazirlanis_secici.get_date()
        body_text2=f"{date_ist:%d.%m.%Y}"
        to_chp=chp_alani.get()
#line sp#ace
        for _ in range(1):
            linespace_style=doc.styles["Body Text"]
            linespace=doc.add_paragraph(style=linespace_style).add_run(" ")
            linespace_style.font.size=Pt(10)

# certificate template

        ist=f"İSTANBUL"
        ist_style = doc.styles['Body Text']
        ist = doc.add_paragraph(style=ist_style).add_run(ist)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ist.font.size = Pt(11)
        ist.font.bold = True
        ist.font.underline =True
        ist.font.name="Arial"

        date_ist_style=doc.styles["Normal"]
        date_ist=doc.add_paragraph(style=date_ist_style).add_run(f"{body_text2}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_ist.font.size = Pt(11)
        date_ist.font.name="Arial"
    


        body_text = f"To: WHOM IT MAY CONCERN;"

        towhom_style = doc.styles['Body Text']
        towhom = doc.add_paragraph(style=towhom_style).add_run(body_text)
        towhom.font.size = Pt(11)
        towhom.font.bold = True
        towhom.font.name="Arial"

        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        heading=f"NON INCIDENT/ACCIDENT STATEMENT"
        heading_style = doc.styles['Body Text']
        head = doc.add_paragraph(style=heading_style).add_run(f'{heading}')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        head.font.size = Pt(11)
        head.font.name="Arial"
        head.font.bold = True 
        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10


        if len(sn_list) > 1:
                
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_textif}")
                body.font.size=Pt(11)
                body.font.name="Arial"
        else:
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_text1}")
                body.font.size=Pt(11)
                body.font.name="Arial"
                body.font.name="Arial"

#TABLO GELİCEK ALAN
        table=doc.add_table(df1.shape[0]+1, df1.shape[1],style="Light List Accent 1")
        table.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.style.font.size=Pt(9)
        
        if len(desc_list)==1 and len(pn_list)==1 and len(sn_list)==1:
                
                for k in range(df1.shape[-1]):
                        table.cell(0,k).text=df1.columns[k]
                for l in range(df1.shape[0]):
                        for k in range(df1.shape[-1]):
                                table.cell(l+1,k).text=str(df1.values[l,k])
                                
        elif len(desc_list)==2 and len(pn_list)==2 and len(sn_list)==2:
                if desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==3 and len(pn_list)==3 and len(sn_list)==3:
                if desc_list[0]==desc_list[2] and pn_list[0]==pn_list[2] and sn_list[0]==sn_list[2]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
                elif desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
                
                
                
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
                
        elif len(desc_list)==4 and len(pn_list)==4 and len(sn_list)==4:
                if desc_list[0]==desc_list[3] and pn_list[0]==pn_list[3] and sn_list[0]==sn_list[3]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[4]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==5 and len(pn_list)==5 and len(sn_list)==5:
                if desc_list[0]==desc_list[4] and pn_list[3]==pn_list[4] and sn_list[0]==sn_list[4]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[5]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif sn_list[0]==sn_list[2] and sn_list[3]==sn_list[4]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[4]
                        rowd=table.rows[5]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
        elif len(desc_list)==6 and len(pn_list)==6 and len(sn_list)==6:
                if desc_list[0]==desc_list[5] and pn_list[0]==pn_list[5] and sn_list[0]==sn_list[5]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[5]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])


                elif sn_list[0]==sn_list[2] and sn_list[3]==sn_list[5]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[4]
                        rowd=table.rows[6]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
        
               
        

        #if len(sn_list) > 1:
        #        
        #        parag_style=doc.styles["Body Text"]
        #        parag=doc.add_paragraph(style=parag_style).add_run(f"{paragraph1}")
        #        parag.font.size=Pt(11)
        #        parag.font.name="Arial"
        #else:
        #        parag_style=doc.styles["Body Text"]
        #        parag=doc.add_paragraph(style=parag_style).add_run(f"{paragraph2}")
        #        parag.font.size=Pt(11)
        #        parag.font.name="Arial"           


        for _ in range(2):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        tai_text=f"TURKISH AIRLINES INC."
        tai_style=doc.styles["Body Text"]
        tai=doc.add_paragraph(style=tai_style).add_run(f"{tai_text}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tai.font.size=Pt(12)
        tai.font.name="Arial"
        tai.font.bold= True

        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        chp_style=doc.styles["Body Text"]
        chp=doc.add_paragraph(style=chp_style).add_run(f"{sign}")
        chp.font.size = Pt(11)
        chp.font.name="Arial"
        chp.font.bold= True
        for _ in range(14):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        
       
            
        
        footer_section=doc.add_picture("footer1.png", width=Inches(5.00),height=Inches(0.61))
        footer_thy=doc.paragraphs[-1]
        footer_thy.alignment=WD_ALIGN_PARAGRAPH.CENTER
       
           
    
        doc.save("word.docx")
        
    elif nis_tipi_opsiyon.get()==var4 and nis_operator_opsiyon.get()==var5:

        doc=Document()
        sections=doc.sections
        



        for section in sections:
            section.top_margin=Inches(0.18)
            section.bottom_margin=Inches(0.38)
            section.left_margin=Inches(0.38)
            section.right_margin=Inches(0.58)

        section=doc.sections[0]



        logo=doc.add_picture("turkishairlines.png", width=Inches(1.97), height=Inches(0.64))
        logo_placement=doc.paragraphs[-1]
        logo_placement.alignment=WD_ALIGN_PARAGRAPH.CENTER


        to_date=tod.get()
        #new_to_date=to_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        #,"August").replace(".09.","September").replace(".10.","October").replace(".11."
        #,"November").replace(".12.","December")
        from_date=fd.get()
        new_from_date=from_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        ztop_date=ztop.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        body_text1=f"      This is to certify that, to the best of our knowledge, the part listed below has not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or was not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of THY operations from {ztop_date} to {new_from_date}."
        body_textif=f"      This is to certify that, to the best of our knowledge, the parts listed below have not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or were not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of THY operations from {ztop_date} to {new_from_date}."
        date_ist=nis_hazirlanis_secici.get_date()
        body_text2=f"{date_ist:%d.%m.%Y}"
        to_chp=chp_alani.get()

        #body_text1=f"      This is to certify that, to the best of our knowledge, the part listed below has not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or was not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of THY operations."
        #body_textif=f"      This is to certify that, to the best of our knowledge, the parts listed below have not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or were not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of THY operations."
        #date_ist=nis_hazirlanis_secici.get_date()
        #body_text2=f"{date_ist:%d.%m.%Y}"
        #to_chp=chp_alani.get()
        
#line sp#ace
        for _ in range(1):
            linespace_style=doc.styles["Body Text"]
            linespace=doc.add_paragraph(style=linespace_style).add_run(" ")
            linespace_style.font.size=Pt(10)


        ist=f"İSTANBUL"
        ist_style = doc.styles['Body Text']
        ist = doc.add_paragraph(style=ist_style).add_run(ist)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ist.font.size = Pt(11)
        ist.font.bold = True
        ist.font.underline =True
        ist.font.name="Arial"

        date_ist_style=doc.styles["Normal"]
        date_ist=doc.add_paragraph(style=date_ist_style).add_run(f"{body_text2}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_ist.font.size = Pt(11)
        date_ist.font.name="Arial"
    


        body_text = f"To: WHOM IT MAY CONCERN;"

        towhom_style = doc.styles['Body Text']
        towhom = doc.add_paragraph(style=towhom_style).add_run(body_text)
        towhom.font.size = Pt(11)
        towhom.font.bold = True
        towhom.font.name="Arial"

        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        heading=f"NON INCIDENT/ACCIDENT STATEMENT"
        heading_style = doc.styles['Body Text']
        head = doc.add_paragraph(style=heading_style).add_run(f'{heading}')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        head.font.size = Pt(11)
        head.font.name="Arial"
        head.font.bold = True 
        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10


        if len(sn_list) > 1:
                
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_textif}")
                body.font.size=Pt(11)
                body.font.name="Arial"
        else:
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_text1}")
                body.font.size=Pt(11)
                body.font.name="Arial"

#TABLO GELİCEK ALAN
        table=doc.add_table(df1.shape[0]+1, df1.shape[1],style="Light List Accent 1")
        table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.style.font.size=Pt(9)
        if len(desc_list)==1 and len(pn_list)==1 and len(sn_list)==1:
                
                for k in range(df1.shape[-1]):
                        table.cell(0,k).text=df1.columns[k]
                for l in range(df1.shape[0]):
                        for k in range(df1.shape[-1]):
                                table.cell(l+1,k).text=str(df1.values[l,k])
                                
        elif len(desc_list)==2 and len(pn_list)==2 and len(sn_list)==2:
                if desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==3 and len(pn_list)==3 and len(sn_list)==3:
                if desc_list[0]==desc_list[2] and pn_list[0]==pn_list[2] and sn_list[0]==sn_list[2]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
                
                
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
                
        elif len(desc_list)==4 and len(pn_list)==4 and len(sn_list)==4:
                if desc_list[0]==desc_list[3] and pn_list[0]==pn_list[3] and sn_list[0]==sn_list[3]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[4]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                elif  pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1] and pn_list[2]==pn_list[3] and sn_list[2]==sn_list[3]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==5 and len(pn_list)==5 and len(sn_list)==5:
                if desc_list[0]==desc_list[4] and pn_list[3]==pn_list[4] and sn_list[0]==sn_list[4]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[5]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif sn_list[0]==sn_list[2] and sn_list[3]==sn_list[4]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[4]
                        rowd=table.rows[5]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==6 and len(pn_list)==6 and len(sn_list)==6:
                if desc_list[0]==desc_list[5] and pn_list[0]==pn_list[5] and sn_list[0]==sn_list[5]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[6]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
        else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])


        for _ in range(2):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        tai_text=f"TURKISH AIRLINES INC."
        tai_style=doc.styles["Body Text"]
        tai=doc.add_paragraph(style=tai_style).add_run(f"{tai_text}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tai.font.size=Pt(12)
        tai.font.name="Arial"
        tai.font.bold= True

        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        chp_style=doc.styles["Body Text"]
        chp=doc.add_paragraph(style=chp_style).add_run(f"{sign}")
        chp.font.size = Pt(11)
        chp.font.name="Arial"
        chp.font.bold= True
        for _ in range(14):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        
        
        footer_section=doc.add_picture("footer1.png", width=Inches(5.00),height=Inches(0.61))
        footer_thy=doc.paragraphs[-1]
        footer_thy.alignment= WD_ALIGN_PARAGRAPH.CENTER
        

        doc.save("word.docx")
    

    elif nis_tipi_opsiyon.get()==var3 and nis_operator_opsiyon.get()==var5:

        doc=Document()
        sections=doc.sections

        for section in sections:
            section.top_margin=Inches(0.18)
            section.bottom_margin=Inches(0.38)
            section.left_margin=Inches(0.38)
            section.right_margin=Inches(0.58)

        section=doc.sections[0]



        logo=doc.add_picture("turkishairlines.png", width=Inches(1.97), height=Inches(0.64))
        logo_placement=doc.paragraphs[-1]
        logo_placement.alignment=WD_ALIGN_PARAGRAPH.CENTER
    
        to_date=tod.get()
        #new_to_date=to_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        #,"August").replace(".09.","September").replace(".10.","October").replace(".11."
        #,"November").replace(".12.","December")
        from_date=fd.get()
        new_from_date=from_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        ztop_date=ztop.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        body_text1=f"      This is to certify that, the part, details of which are specified below,has been operated by Turkish Airlines during the period from {ztop_date} to {new_from_date}."
        body_textif=f"      This is to certify that, the parts, details of which are specified below, has been operated by Turkish Airlines during the period from {ztop_date} to {new_from_date}."
        paragraph1=f"      I hereby certify that, to the best of my knowledge, during the period stated above:\n\t1. No subject part installed has been;\n\n\t\ta. damaged during, or identified as the root cause of, a reportable incident or \t\t\taccident as defined by Annex 13 to the Chicago Convention, other than the \t\t\tfollowing:\n\n\t\t\t-Write The Occurence Report Note\n\n\t\tb.subjected to severe stress or heat such as in a major engine failure, accident,or\n\t\t fire or has been submersed in salt water,\n\n\t\t unless its airworthiness status was re-established by an approved maintenance\n\t\t organisation in accordance with the applicable airworthiness regulations and\n\t\t instructions of the type certificate holder and-or supplemental type certificate\n\t\t holder and-or OEM of the part, and supported by an authorised airworthiness\n\t\t release certificate.\n\n\t2. No part has been installed on the aircraft which was obtained from a military source or\n\t was previously fitted to a state aircraft as deemed by Article 3 of the Chicago\n\t Convention. "
        date_ist=nis_hazirlanis_secici.get_date()
        body_text2=f"{date_ist:%d.%m.%Y}"
        to_chp=chp_alani.get()
        
#line space
        for _ in range(1):
            linespace_style=doc.styles["Body Text"]
            linespace=doc.add_paragraph(style=linespace_style).add_run(" ")
            linespace_style.font.size=Pt(10)


        ist=f"İSTANBUL"
        ist_style = doc.styles['Body Text']
        ist = doc.add_paragraph(style=ist_style).add_run(ist)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ist.font.size = Pt(11)
        ist.font.bold = True
        ist.font.underline =True
        ist.font.name="Arial"

        date_ist_style=doc.styles["Normal"]
        date_ist=doc.add_paragraph(style=date_ist_style).add_run(f"{body_text2}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_ist.font.size = Pt(11)
        date_ist.font.name="Arial"
    


        body_text = f"To: WHOM IT MAY CONCERN;"

        towhom_style = doc.styles['Body Text']
        towhom = doc.add_paragraph(style=towhom_style).add_run(body_text)
        towhom.font.size = Pt(11)
        towhom.font.bold = True
        towhom.font.name="Arial"

        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        heading=f"NON INCIDENT/ACCIDENT STATEMENT"
        heading_style = doc.styles['Body Text']
        head = doc.add_paragraph(style=heading_style).add_run(f'{heading}')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        head.font.size = Pt(11)
        head.font.name="Arial"
        head.font.bold = True 

        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10


        if len(sn_list) > 1:
                
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_textif}")
                body.font.size=Pt(11)
                body.font.name="Arial"
        else:
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_text1}")
                body.font.size=Pt(11)
                body.font.name="Arial"

#TABLO GELİCEK ALAN
        table=doc.add_table(df1.shape[0]+1, df1.shape[1],style="Light List Accent 1")
        table.style.font.size=Pt(9)
        table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.width=Inches(3.1)
        if len(desc_list)==1 and len(pn_list)==1 and len(sn_list)==1:
                
                for k in range(df1.shape[-1]):
                        table.cell(0,k).text=df1.columns[k]
                for l in range(df1.shape[0]):
                        for k in range(df1.shape[-1]):
                                table.cell(l+1,k).text=str(df1.values[l,k])
                                
        elif len(desc_list)==2 and len(pn_list)==2 and len(sn_list)==2:
                if desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==3 and len(pn_list)==3 and len(sn_list)==3:
                if desc_list[0]==desc_list[2] and pn_list[0]==pn_list[2] and sn_list[0]==sn_list[2]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
                
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
                
        elif len(desc_list)==4 and len(pn_list)==4 and len(sn_list)==4:
                if desc_list[0]==desc_list[3] and pn_list[0]==pn_list[3] and sn_list[0]==sn_list[3]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[4]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                elif  pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1] and pn_list[2]==pn_list[3] and sn_list[2]==sn_list[3]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==5 and len(pn_list)==5 and len(sn_list)==5:
                if desc_list[0]==desc_list[4] and pn_list[3]==pn_list[4] and sn_list[0]==sn_list[4]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[5]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif sn_list[0]==sn_list[2] and sn_list[3]==sn_list[4]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[4]
                        rowd=table.rows[5]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==6 and len(pn_list)==6 and len(sn_list)==6:
                if desc_list[0]==desc_list[5] and pn_list[0]==pn_list[5] and sn_list[0]==sn_list[5]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[6]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif sn_list[0]==sn_list[2] and sn_list[3]==sn_list[5]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[4]
                        rowd=table.rows[6]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])


                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])


                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

        elif len(desc_list)==7 and len(pn_list)==7 and len(sn_list)==7:
                if desc_list[0]==desc_list[6] and pn_list[0]==pn_list[6] and sn_list[0]==sn_list[6]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[7]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])


                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==8 and len(pn_list)==8 and len(sn_list)==8:
                if desc_list[0]==desc_list[7] and pn_list[0]==pn_list[7] and sn_list[0]==sn_list[7]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[8]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==9 and len(pn_list)==9 and len(sn_list)==9:
                if desc_list[0]==desc_list[8] and pn_list[0]==pn_list[8] and sn_list[0]==sn_list[8]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[9]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
        elif len(desc_list)==10 and len(pn_list)==10 and len(sn_list)==10:
                if desc_list[0]==desc_list[9] and pn_list[0]==pn_list[9] and sn_list[0]==sn_list[9]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[10]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                
                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5] and sn_list[6]==sn_list[7] and sn_list[8]==sn_list[9]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])
                        rowg=table.rows[7]
                        rowh=table.rows[8]
                        rowg.cells[0].merge(rowh.cells[0])
                        rowg.cells[1].merge(rowh.cells[1])
                        rowg.cells[2].merge(rowh.cells[2])
                        rowj=table.rows[9]
                        rowk=table.rows[10]
                        rowj.cells[0].merge(rowk.cells[0])
                        rowj.cells[1].merge(rowk.cells[1])
                        rowj.cells[2].merge(rowk.cells[2])
                        
                        
                        
                        
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==11 and len(pn_list)==11 and len(sn_list)==11:
                if desc_list[0]==desc_list[10] and pn_list[0]==pn_list[10] and sn_list[0]==sn_list[10]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[11]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==12 and len(pn_list)==12 and len(sn_list)==12:
                if desc_list[0]==desc_list[11] and pn_list[0]==pn_list[11] and sn_list[0]==sn_list[11]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[12]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
        else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])         

                
                        
        parag_style=doc.styles["Body Text"]
        parag=doc.add_paragraph(style=parag_style).add_run(f"{paragraph1}")
        parag.font.size=Pt(11)
        parag.font.name="Arial"
        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        tai_text=f"TURKISH AIRLINES INC."
        tai_style=doc.styles["Body Text"]
        tai=doc.add_paragraph(style=tai_style).add_run(f"{tai_text}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tai.font.size=Pt(12)
        tai.font.name="Arial"
        tai.font.bold= True

        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        chp_style=doc.styles["Body Text"]
        chp=doc.add_paragraph(style=chp_style).add_run(f"{sign}")
        chp.font.size = Pt(11)
        chp.font.name="Arial"
        chp.font.bold= True

        for _ in range(14):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10
        
        footer_section=doc.add_picture("footer1.png", width=Inches(5.00),height=Inches(0.61))
        footer_thy=doc.paragraphs[-1]
        footer_thy.alignment= WD_ALIGN_PARAGRAPH.CENTER
        

        doc.save("word.docx")
    #AJET NIS ŞABLONU DOCX
    elif nis_tipi_opsiyon.get()==var1 and nis_operator_opsiyon.get()==var6:
            
        

        doc=Document()
        sections=doc.sections


        for section in sections:
            section.top_margin=Inches(0.18)
            section.bottom_margin=Inches(0.38)
            section.left_margin=Inches(0.38)
            section.right_margin=Inches(0.58)

        section=doc.sections[0]
        logo=doc.add_picture("ajet.png", width=Inches(1.97), height=Inches(0.64))
        logo_placement=doc.paragraphs[-1]
        logo_placement.alignment=WD_ALIGN_PARAGRAPH.CENTER
    
        #to_date=tod.get()
        #new_to_date=to_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        #,"August").replace(".09.","September").replace(".10.","October").replace(".11."
        #,"November").replace(".12.","December")
        #from_date=fd.get()
        #new_from_date=from_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        #,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        #body_text1=f"      This is to certify that, to the best of our knowledge, the part listed below has not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or was not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of THY operations."
        #body_textif=f"      This is to certify that, to the best of our knowledge, the parts listed below have not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or were not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of THY operations."
        #date_ist=nis_hazirlanis_secici.get_date()
        #body_text2=f"{date_ist:%d.%m.%Y}"
        #to_chp=chp_alani.get()
        



        to_date=tod.get()
        #new_to_date=to_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        #,"August").replace(".09.","September").replace(".10.","October").replace(".11."
        #,"November").replace(".12.","December")
        from_date=fd.get()
        new_from_date=from_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        ztop_date=ztop.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        body_text1=f"      This is to certify that, to the best of our knowledge, the part listed below has not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or was not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of AJET operations from {ztop_date} to {new_from_date}."
        body_textif=f"      This is to certify that, to the best of our knowledge, the parts listed below have not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or were not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of AJET operations from {ztop_date} to {new_from_date}."
        paragraph1=f"      I hereby certify that, to the best of my knowledge, during the period stated above:\n\t1. No subject part installed has been;\n\n\t\ta. damaged during, or identified as the root cause of, a reportable incident or \t\t\taccident as defined by Annex 13 to the Chicago Convention, other than the \t\t\tfollowing:\n\n\t\t\t-Write The Occurence Report Note\n\n\t\tb.subjected to severe stress or heat such as in a major engine failure, accident,or\n\t\t fire or has been submersed in salt water,\n\n\t\t unless its airworthiness status was re-established by an approved maintenance\n\t\t organisation in accordance with the applicable airworthiness regulations and\n\t\t instructions of the type certificate holder and-or supplemental type certificate\n\t\t holder and-or OEM of the part, and supported by an authorised airworthiness\n\t\t release certificate.\n\n\t2. No part has been installed on the aircraft which was obtained from a military source or\n\t was previously fitted to a state aircraft as deemed by Article 3 of the Chicago\n\t Convention. "
        date_ist=nis_hazirlanis_secici.get_date()
        body_text2=f"{date_ist:%d.%m.%Y}"
        to_chp=chp_alani.get()
#line sp#ace
        for _ in range(1):
            linespace_style=doc.styles["Body Text"]
            linespace=doc.add_paragraph(style=linespace_style).add_run(" ")
            linespace_style.font.size=Pt(10)

# certificate template

        ist=f"İSTANBUL"
        ist_style = doc.styles['Body Text']
        ist = doc.add_paragraph(style=ist_style).add_run(ist)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ist.font.size = Pt(11)
        ist.font.bold = True
        ist.font.underline =True
        ist.font.name="Arial"

        date_ist_style=doc.styles["Normal"]
        date_ist=doc.add_paragraph(style=date_ist_style).add_run(f"{body_text2}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_ist.font.size = Pt(11)
        date_ist.font.name="Arial"
    


        body_text = f"To: WHOM IT MAY CONCERN;"

        towhom_style = doc.styles['Body Text']
        towhom = doc.add_paragraph(style=towhom_style).add_run(body_text)
        towhom.font.size = Pt(11)
        towhom.font.bold = True
        towhom.font.name="Arial"

        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        heading=f"NON INCIDENT/ACCIDENT STATEMENT"
        heading_style = doc.styles['Body Text']
        head = doc.add_paragraph(style=heading_style).add_run(f'{heading}')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        head.font.size = Pt(11)
        head.font.name="Arial"
        head.font.bold = True 
        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10


        if len(sn_list) > 1:
                
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_textif}")
                body.font.size=Pt(11)
                body.font.name="Arial"
        else:
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_text1}")
                body.font.size=Pt(11)
                body.font.name="Arial"
                body.font.name="Arial"

#TABLO GELİCEK ALAN
        table=doc.add_table(df1.shape[0]+1, df1.shape[1],style="Light List Accent 1")
        table.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.style.font.size=Pt(9)
        
        if len(desc_list)==1 and len(pn_list)==1 and len(sn_list)==1:
                
                for k in range(df1.shape[-1]):
                        table.cell(0,k).text=df1.columns[k]
                for l in range(df1.shape[0]):
                        for k in range(df1.shape[-1]):
                                table.cell(l+1,k).text=str(df1.values[l,k])
                                
        elif len(desc_list)==2 and len(pn_list)==2 and len(sn_list)==2:
                if desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==3 and len(pn_list)==3 and len(sn_list)==3:
                if desc_list[0]==desc_list[2] and pn_list[0]==pn_list[2] and sn_list[0]==sn_list[2]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
                elif desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
                
                
                
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
                
        elif len(desc_list)==4 and len(pn_list)==4 and len(sn_list)==4:
                if desc_list[0]==desc_list[3] and pn_list[0]==pn_list[3] and sn_list[0]==sn_list[3]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[4]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==5 and len(pn_list)==5 and len(sn_list)==5:
                if desc_list[0]==desc_list[4] and pn_list[3]==pn_list[4] and sn_list[0]==sn_list[4]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[5]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif sn_list[0]==sn_list[2] and sn_list[3]==sn_list[4]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[4]
                        rowd=table.rows[5]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
        elif len(desc_list)==6 and len(pn_list)==6 and len(sn_list)==6:
                if desc_list[0]==desc_list[5] and pn_list[0]==pn_list[5] and sn_list[0]==sn_list[5]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[5]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])


                elif sn_list[0]==sn_list[2] and sn_list[3]==sn_list[5]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[4]
                        rowd=table.rows[6]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
        
               
        

        #if len(sn_list) > 1:
        #        
        #        parag_style=doc.styles["Body Text"]
        #        parag=doc.add_paragraph(style=parag_style).add_run(f"{paragraph1}")
        #        parag.font.size=Pt(11)
        #        parag.font.name="Arial"
        #else:
        #        parag_style=doc.styles["Body Text"]
        #        parag=doc.add_paragraph(style=parag_style).add_run(f"{paragraph2}")
        #        parag.font.size=Pt(11)
        #        parag.font.name="Arial"           


        for _ in range(2):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10
        chp_style=doc.styles["Body Text"]
        chp=doc.add_paragraph(style=chp_style).add_run(f"{sign}")
        chp.font.size = Pt(11)
        chp.font.name="Arial"
        chp.font.bold= True

        for _ in range(9):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10
        tai_text=f"AJet Hava Taşımacılığı Anonim Şirketi\n        Yeşilköy Mah. Havaalanı Cad. THY Genel Müdürlük No:3/1\n        Bakırköy / İstanbul "
        tai_style=doc.styles["Body Text"]
        tai=doc.add_paragraph(style=tai_style).add_run(f"{tai_text}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tai.font.size=Pt(8)
        tai.font.name="Arial"
        tai.font.bold= True

        doc.save("word.docx")
    if nis_tipi_opsiyon.get()==var2 and nis_operator_opsiyon.get()==var6:
            
        
        doc=Document()
        sections=doc.sections
        
        for section in sections:
            section.top_margin=Inches(0.18)
            section.bottom_margin=Inches(0.38)
            section.left_margin=Inches(0.38)
            section.right_margin=Inches(0.58)

        section=doc.sections[0]



        logo=doc.add_picture("ajet.png", width=Inches(1.97), height=Inches(0.64))
        logo_placement=doc.paragraphs[-1]
        logo_placement.alignment=WD_ALIGN_PARAGRAPH.CENTER
        
        to_date=tod.get()
        #new_to_date=to_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        #,"August").replace(".09.","September").replace(".10.","October").replace(".11."
        #,"November").replace(".12.","December")
        from_date=fd.get()
        new_from_date=from_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        ztop_date=ztop.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        body_text1=f"      This is to certify that, to the best of our knowledge, the part listed below has not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or was not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of AJET operations from {ztop_date} to {new_from_date}."
        body_textif=f"      This is to certify that, to the best of our knowledge, the parts listed below have not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or were not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of AJET operations from {ztop_date} to {new_from_date}."
        date_ist=nis_hazirlanis_secici.get_date()
        body_text2=f"{date_ist:%d.%m.%Y}"
        to_chp=chp_alani.get()
        
        
#line sp#ace
        for _ in range(1):
            linespace_style=doc.styles["Body Text"]
            linespace=doc.add_paragraph(style=linespace_style).add_run(" ")
            linespace_style.font.size=Pt(10)
#
# certificate template

        ist=f"İSTANBUL"
        ist_style = doc.styles['Body Text']
        ist = doc.add_paragraph(style=ist_style).add_run(ist)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ist.font.size = Pt(11)
        ist.font.bold = True
        ist.font.underline =True
        ist.font.name="Arial"

        date_ist_style=doc.styles["Normal"]
        date_ist=doc.add_paragraph(style=date_ist_style).add_run(f"{body_text2}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_ist.font.size = Pt(11)
        date_ist.font.name="Arial"
    


        body_text = f"To: WHOM IT MAY CONCERN;"

        towhom_style = doc.styles['Body Text']
        towhom = doc.add_paragraph(style=towhom_style).add_run(body_text)
        towhom.font.size = Pt(11)
        towhom.font.bold = True
        towhom.font.name="Arial"

        #for _ in range(1):
                #linespace_style = doc.styles['Body Text']
                #linespace = doc.add_paragraph(style=linespace_style).add_run()
                #linespace.font.size = 10

        heading=f"NON INCIDENT/ACCIDENT STATEMENT"
        heading_style = doc.styles['Body Text']
        head = doc.add_paragraph(style=heading_style).add_run(f'{heading}')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        head.font.size = Pt(11)
        head.font.name="Arial"
        head.font.bold = True 
        #for _ in range(1):
                #linespace_style = doc.styles['Body Text']
                #linespace = doc.add_paragraph(style=linespace_style).add_run()
                #linespace.font.size = 10

        if len(sn_list) > 1:
                
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_textif}")
                body.font.size=Pt(11)
                body.font.name="Arial"
        else:
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_text1}")
                body.font.size=Pt(11)
                body.font.name="Arial"
#Table
        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10
        table=doc.add_table(df1.shape[0]+1, df1.shape[1],style="Light List Accent 1")
        table.style.font.size=Pt(9)
        table.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.border=True
        #def ChangeWidthOfTable(table,width,column):
        #        for columnVariable in range(0,column):
        #                for cell in table.columns[columnVariable].cells:
        #                        cell.width=width(15)
        #red="TC-JFN 29776"
        #if ac_list[0]==red:
        #        red.font.color.rgb=RGBColor(255,0,0)
        #        print(red)
        

                #r="\033[1;32m] TC-JFN 29776"
                #ac_list[0]="\033[31mTC-JFN 29776\033[0m"
                #print(ac_list[0]) 
        if len(desc_list)==1 and len(pn_list)==1 and len(sn_list)==1:
                
                for k in range(df1.shape[-1]):
                        table.cell(0,k).text=df1.columns[k]
                for l in range(df1.shape[0]):
                        for k in range(df1.shape[-1]):
                                table.cell(l+1,k).text=str(df1.values[l,k])
                                
        elif len(desc_list)==2 and len(pn_list)==2 and len(sn_list)==2:
                if desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==3 and len(pn_list)==3 and len(sn_list)==3:
                if desc_list[0]==desc_list[2] and pn_list[0]==pn_list[2] and sn_list[0]==sn_list[2]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
                elif desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
                elif sn_list[1]==sn_list[2] and pn_list[1]==pn_list[2]:
                        rowa=table.rows[2]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
                
        elif len(desc_list)==4 and len(pn_list)==4 and len(sn_list)==4:
                if desc_list[0]==desc_list[3] and pn_list[0]==pn_list[3] and sn_list[0]==sn_list[3]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[4]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                elif  pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1] and pn_list[2]==pn_list[3] and sn_list[2]==sn_list[3]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
                elif desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                elif sn_list[0]==sn_list[2] and sn_list[2]!=sn_list[3]:
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])


                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==5 and len(pn_list)==5 and len(sn_list)==5:
                if desc_list[0]==desc_list[4] and pn_list[0]==pn_list[4] and sn_list[0]==sn_list[4]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[5]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])


                elif sn_list[0]==sn_list[2] and sn_list[3]==sn_list[4]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[4]
                        rowd=table.rows[5]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
        elif len(desc_list)==6 and len(pn_list)==6 and len(sn_list)==6:
                if desc_list[0]==desc_list[5] and pn_list[0]==pn_list[5] and sn_list[0]==sn_list[5]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[6]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                

                elif sn_list[0]==sn_list[2] and sn_list[3]==sn_list[5]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowc=table.rows[3]
                        #rowd=table.rows[5]
                        #rowc.cells[0].merge(rowd.cells[0])
                        #rowc.cells[1].merge(rowd.cells[1])
                        #rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[4]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])


                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])



                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==7 and len(pn_list)==7 and len(sn_list)==7:
                if desc_list[0]==desc_list[6] and pn_list[0]==pn_list[6] and sn_list[0]==sn_list[6]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[7]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])


                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==8 and len(pn_list)==8 and len(sn_list)==8:
                if desc_list[0]==desc_list[7] and pn_list[0]==pn_list[7] and sn_list[0]==sn_list[7]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[8]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[3]!=sn_list[4] and  sn_list[4]==sn_list[5] and sn_list[6]==sn_list[7]:
                        rowa=table.rows[1]
                        rowb=table.rows[4]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[5]
                        rowd=table.rows[8]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])


                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5] and sn_list[6]==sn_list[7]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])
                        rowg=table.rows[7]
                        rowh=table.rows[8]
                        rowg.cells[0].merge(rowh.cells[0])
                        rowg.cells[1].merge(rowh.cells[1])
                        rowg.cells[2].merge(rowh.cells[2])
                        

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==9 and len(pn_list)==9 and len(sn_list)==9:
                if desc_list[0]==desc_list[8] and pn_list[0]==pn_list[8] and sn_list[0]==sn_list[8]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[9]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
        elif len(desc_list)==10 and len(pn_list)==10 and len(sn_list)==10:
                if desc_list[0]==desc_list[9] and pn_list[0]==pn_list[9] and sn_list[0]==sn_list[9]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[10]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                
                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5] and sn_list[6]==sn_list[7] and sn_list[8]==sn_list[9]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])
                        rowg=table.rows[7]
                        rowh=table.rows[8]
                        rowg.cells[0].merge(rowh.cells[0])
                        rowg.cells[1].merge(rowh.cells[1])
                        rowg.cells[2].merge(rowh.cells[2])
                        rowj=table.rows[9]
                        rowk=table.rows[10]
                        rowj.cells[0].merge(rowk.cells[0])
                        rowj.cells[1].merge(rowk.cells[1])
                        rowj.cells[2].merge(rowk.cells[2])
                        
                        
                        
                        
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==11 and len(pn_list)==11 and len(sn_list)==11:
                if desc_list[0]==desc_list[10] and pn_list[0]==pn_list[10] and sn_list[0]==sn_list[10]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[11]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==12 and len(pn_list)==12 and len(sn_list)==12:
                if desc_list[0]==desc_list[11] and pn_list[0]==pn_list[11] and sn_list[0]==sn_list[11]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[12]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
        else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])                

        for _ in range(2):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10
         

        
        chp_style=doc.styles["Body Text"]
        chp=doc.add_paragraph(style=chp_style).add_run(f"{sign}")
        chp.font.size = Pt(11)
        chp.font.name="Arial"
        chp.font.bold= True
        

        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        

        for _ in range(10):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10
        
        tai_text=f"AJet Hava Taşımacılığı Anonim Şirketi\n        Yeşilköy Mah. Havaalanı Cad. THY Genel Müdürlük No:3/1\n        Bakırköy / İstanbul "
        tai_style=doc.styles["Body Text"]
        tai=doc.add_paragraph(style=tai_style).add_run(f"{tai_text}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tai.font.size=Pt(8)
        tai.font.name="Arial"
        tai.font.bold= True

    
    
        doc.save("word.docx")
    elif nis_tipi_opsiyon.get()==var3 and nis_operator_opsiyon.get()==var6:
        doc=Document()
        sections=doc.sections

        for section in sections:
            section.top_margin=Inches(0.18)
            section.bottom_margin=Inches(0.38)
            section.left_margin=Inches(0.38)
            section.right_margin=Inches(0.58)

        section=doc.sections[0]



        logo=doc.add_picture("ajet.png", width=Inches(1.97), height=Inches(0.64))
        logo_placement=doc.paragraphs[-1]
        logo_placement.alignment=WD_ALIGN_PARAGRAPH.CENTER
    
        to_date=tod.get()
        #new_to_date=to_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        #,"August").replace(".09.","September").replace(".10.","October").replace(".11."
        #,"November").replace(".12.","December")
        from_date=fd.get()
        new_from_date=from_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        ztop_date=ztop.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        body_text1=f"      This is to certify that, the part, details of which are specified below,has been operated by Ajet during the period from {ztop_date} to {new_from_date}."
        body_textif=f"      This is to certify that, the parts, details of which are specified below, has been operated by Ajet during the period from {ztop_date} to {new_from_date}."
        paragraph1=f"      I hereby certify that, to the best of my knowledge, during the period stated above:\n\t1. No subject parts installed have been;\n\n\t\ta. damaged during, or identified as the root cause of, a reportable incident or \t\t\taccident as defined by Annex 13 to the Chicago Convention, other than the \t\t\tfollowing:\n\n\t\t\t-Write The Occurence Report Note\n\n\t\tb.subjected to severe stress or heat such as in a major engine failure, accident,or\n\t\t fire or has been submersed in salt water,\n\n\t\t unless its airworthiness status was re-established by an approved maintenance\n\t\t organisation in accordance with the applicable airworthiness regulations and\n\t\t instructions of the type certificate holder and-or supplemental type certificate\n\t\t holder and-or OEM of the part, and supported by an authorised airworthiness\n\t\t release certificate.\n\n\t2. No part has been installed on the aircraft which was obtained from a military source or\n\t was previously fitted to a state aircraft as deemed by Article 3 of the Chicago\n\t Convention. "
        paragraph2=f"      I hereby certify that, to the best of my knowledge, during the period stated above:\n\t1. No subject part installed has been;\n\n\t\ta. damaged during, or identified as the root cause of, a reportable incident or \t\t\taccident as defined by Annex 13 to the Chicago Convention, other than the \t\t\tfollowing:\n\n\t\t\t-Write The Occurence Report Note\n\n\t\tb.subjected to severe stress or heat such as in a major engine failure, accident,or\n\t\t fire or has been submersed in salt water,\n\n\t\t unless its airworthiness status was re-established by an approved maintenance\n\t\t organisation in accordance with the applicable airworthiness regulations and\n\t\t instructions of the type certificate holder and-or supplemental type certificate\n\t\t holder and-or OEM of the part, and supported by an authorised airworthiness\n\t\t release certificate.\n\n\t2. No part has been installed on the aircraft which was obtained from a military source or\n\t was previously fitted to a state aircraft as deemed by Article 3 of the Chicago\n\t Convention. "
        date_ist=nis_hazirlanis_secici.get_date()
        body_text2=f"{date_ist:%d.%m.%Y}"
        to_chp=chp_alani.get()
        
#line space
        for _ in range(1):
            linespace_style=doc.styles["Body Text"]
            linespace=doc.add_paragraph(style=linespace_style).add_run(" ")
            linespace_style.font.size=Pt(10)


        ist=f"İSTANBUL"
        ist_style = doc.styles['Body Text']
        ist = doc.add_paragraph(style=ist_style).add_run(ist)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ist.font.size = Pt(11)
        ist.font.bold = True
        ist.font.underline =True
        ist.font.name="Arial"

        date_ist_style=doc.styles["Normal"]
        date_ist=doc.add_paragraph(style=date_ist_style).add_run(f"{body_text2}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_ist.font.size = Pt(11)
        date_ist.font.name="Arial"
    


        body_text = f"To: WHOM IT MAY CONCERN;"

        towhom_style = doc.styles['Body Text']
        towhom = doc.add_paragraph(style=towhom_style).add_run(body_text)
        towhom.font.size = Pt(11)
        towhom.font.bold = True
        towhom.font.name="Arial"

        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        heading=f"NON INCIDENT/ACCIDENT STATEMENT"
        heading_style = doc.styles['Body Text']
        head = doc.add_paragraph(style=heading_style).add_run(f'{heading}')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        head.font.size = Pt(11)
        head.font.name="Arial"
        head.font.bold = True 

        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10


        if len(sn_list) > 1:
                
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_textif}")
                body.font.size=Pt(11)
                body.font.name="Arial"
        else:
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_text1}")
                body.font.size=Pt(11)
                body.font.name="Arial"

#TABLO GELİCEK ALAN
        table=doc.add_table(df1.shape[0]+1, df1.shape[1],style="Light List Accent 1")
        table.style.font.size=Pt(9)
        table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.width=Inches(3.1)
        if len(desc_list)==1 and len(pn_list)==1 and len(sn_list)==1:
                
                for k in range(df1.shape[-1]):
                        table.cell(0,k).text=df1.columns[k]
                for l in range(df1.shape[0]):
                        for k in range(df1.shape[-1]):
                                table.cell(l+1,k).text=str(df1.values[l,k])
                                
        elif len(desc_list)==2 and len(pn_list)==2 and len(sn_list)==2:
                if desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==3 and len(pn_list)==3 and len(sn_list)==3:
                if desc_list[0]==desc_list[2] and pn_list[0]==pn_list[2] and sn_list[0]==sn_list[2]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
                
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
                
        elif len(desc_list)==4 and len(pn_list)==4 and len(sn_list)==4:
                if desc_list[0]==desc_list[3] and pn_list[0]==pn_list[3] and sn_list[0]==sn_list[3]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[4]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                elif  pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1] and pn_list[2]==pn_list[3] and sn_list[2]==sn_list[3]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==5 and len(pn_list)==5 and len(sn_list)==5:
                if desc_list[0]==desc_list[4] and pn_list[3]==pn_list[4] and sn_list[0]==sn_list[4]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[5]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif sn_list[0]==sn_list[2] and sn_list[3]==sn_list[4]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[4]
                        rowd=table.rows[5]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==6 and len(pn_list)==6 and len(sn_list)==6:
                if desc_list[0]==desc_list[5] and pn_list[0]==pn_list[5] and sn_list[0]==sn_list[5]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[6]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif sn_list[0]==sn_list[2] and sn_list[3]==sn_list[5]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[4]
                        rowd=table.rows[6]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])


                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])


                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

        elif len(desc_list)==7 and len(pn_list)==7 and len(sn_list)==7:
                if desc_list[0]==desc_list[6] and pn_list[0]==pn_list[6] and sn_list[0]==sn_list[6]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[7]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])


                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==8 and len(pn_list)==8 and len(sn_list)==8:
                if desc_list[0]==desc_list[7] and pn_list[0]==pn_list[7] and sn_list[0]==sn_list[7]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[8]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==9 and len(pn_list)==9 and len(sn_list)==9:
                if desc_list[0]==desc_list[8] and pn_list[0]==pn_list[8] and sn_list[0]==sn_list[8]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[9]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
        elif len(desc_list)==10 and len(pn_list)==10 and len(sn_list)==10:
                if desc_list[0]==desc_list[9] and pn_list[0]==pn_list[9] and sn_list[0]==sn_list[9]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[10]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                
                elif sn_list[0]==sn_list[1] and sn_list[2]==sn_list[3] and sn_list[4]==sn_list[5] and sn_list[6]==sn_list[7] and sn_list[8]==sn_list[9]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        rowe=table.rows[5]
                        rowf=table.rows[6]
                        rowe.cells[0].merge(rowf.cells[0])
                        rowe.cells[1].merge(rowf.cells[1])
                        rowe.cells[2].merge(rowf.cells[2])
                        rowg=table.rows[7]
                        rowh=table.rows[8]
                        rowg.cells[0].merge(rowh.cells[0])
                        rowg.cells[1].merge(rowh.cells[1])
                        rowg.cells[2].merge(rowh.cells[2])
                        rowj=table.rows[9]
                        rowk=table.rows[10]
                        rowj.cells[0].merge(rowk.cells[0])
                        rowj.cells[1].merge(rowk.cells[1])
                        rowj.cells[2].merge(rowk.cells[2])
                        
                        
                        
                        
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==11 and len(pn_list)==11 and len(sn_list)==11:
                if desc_list[0]==desc_list[10] and pn_list[0]==pn_list[10] and sn_list[0]==sn_list[10]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[11]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==12 and len(pn_list)==12 and len(sn_list)==12:
                if desc_list[0]==desc_list[11] and pn_list[0]==pn_list[11] and sn_list[0]==sn_list[11]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[12]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
        else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])         

                





        if len(sn_list) > 1:
                
                parag_style=doc.styles["Body Text"]
                parag=doc.add_paragraph(style=parag_style).add_run(f"{paragraph1}")
                parag.font.size=Pt(11)
                parag.font.name="Arial"
        else:
                parag_style=doc.styles["Body Text"]
                parag=doc.add_paragraph(style=parag_style).add_run(f"{paragraph2}")
                parag.font.size=Pt(11)
                parag.font.name="Arial"             
        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10
        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        chp_style=doc.styles["Body Text"]
        chp=doc.add_paragraph(style=chp_style).add_run(f"{sign}")
        chp.font.size = Pt(11)
        chp.font.name="Arial"
        chp.font.bold= True

        for _ in range(3):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        tai_text=f"AJet Hava Taşımacılığı Anonim Şirketi\n        Yeşilköy Mah. Havaalanı Cad. THY Genel Müdürlük No:3/1\n        Bakırköy / İstanbul "
        tai_style=doc.styles["Body Text"]
        tai=doc.add_paragraph(style=tai_style).add_run(f"{tai_text}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tai.font.size=Pt(8)
        tai.font.name="Arial"
        tai.font.bold= True

        doc.save("word.docx")
                   
    elif nis_tipi_opsiyon.get()==var4 and nis_operator_opsiyon.get()==var6:

        doc=Document()
        sections=doc.sections
        



        for section in sections:
            section.top_margin=Inches(0.18)
            section.bottom_margin=Inches(0.38)
            section.left_margin=Inches(0.38)
            section.right_margin=Inches(0.58)

        section=doc.sections[0]
        logo=doc.add_picture("ajet.png", width=Inches(1.97), height=Inches(0.64))
        logo_placement=doc.paragraphs[-1]
        logo_placement.alignment=WD_ALIGN_PARAGRAPH.CENTER


        to_date=tod.get()
        #new_to_date=to_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        #,"August").replace(".09.","September").replace(".10.","October").replace(".11."
        #,"November").replace(".12.","December")
        from_date=fd.get()
        new_from_date=from_date.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        ztop_date=ztop.replace(".01.","January").replace(".02.","February").replace(".03.","March").replace(".04.","April").replace(".05.","May").replace(".06.","June").replace(".07.","July").replace(".08."
        ,"August").replace(".09.","September").replace(".10.","October").replace(".11.","November").replace(".12.","December")
        body_text1=f"      This is to certify that, to the best of our knowledge, the part listed below has not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or was not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of AJET operations from {ztop_date} to {new_from_date}."
        body_textif=f"      This is to certify that, to the best of our knowledge, the parts listed below have not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or were not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of AJET operations from {ztop_date} to {new_from_date}."
        date_ist=nis_hazirlanis_secici.get_date()
        body_text2=f"{date_ist:%d.%m.%Y}"
        to_chp=chp_alani.get()

        #body_text1=f"      This is to certify that, to the best of our knowledge, the part listed below has not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or was not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of THY operations."
        #body_textif=f"      This is to certify that, to the best of our knowledge, the parts listed below have not been involved in any incident or accident, nor been subjected to fire or extreme heat or stress or immersed in salt water or were not obtained from any government or military source or subjected to exceedance as defined in ICAO Annex 13 Chapter 1 during the period of THY operations."
        #date_ist=nis_hazirlanis_secici.get_date()
        #body_text2=f"{date_ist:%d.%m.%Y}"
        #to_chp=chp_alani.get()
        
#line sp#ace
        for _ in range(1):
            linespace_style=doc.styles["Body Text"]
            linespace=doc.add_paragraph(style=linespace_style).add_run(" ")
            linespace_style.font.size=Pt(10)


        ist=f"İSTANBUL"
        ist_style = doc.styles['Body Text']
        ist = doc.add_paragraph(style=ist_style).add_run(ist)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ist.font.size = Pt(11)
        ist.font.bold = True
        ist.font.underline =True
        ist.font.name="Arial"

        date_ist_style=doc.styles["Normal"]
        date_ist=doc.add_paragraph(style=date_ist_style).add_run(f"{body_text2}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_ist.font.size = Pt(11)
        date_ist.font.name="Arial"
    


        body_text = f"To: WHOM IT MAY CONCERN;"

        towhom_style = doc.styles['Body Text']
        towhom = doc.add_paragraph(style=towhom_style).add_run(body_text)
        towhom.font.size = Pt(11)
        towhom.font.bold = True
        towhom.font.name="Arial"

        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        heading=f"NON INCIDENT/ACCIDENT STATEMENT"
        heading_style = doc.styles['Body Text']
        head = doc.add_paragraph(style=heading_style).add_run(f'{heading}')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        head.font.size = Pt(11)
        head.font.name="Arial"
        head.font.bold = True 
        for _ in range(1):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10


        if len(sn_list) > 1:
                
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_textif}")
                body.font.size=Pt(11)
                body.font.name="Arial"
        else:
                body_style=doc.styles["Body Text"]
                body=doc.add_paragraph(style=body_style).add_run(f"{body_text1}")
                body.font.size=Pt(11)
                body.font.name="Arial"

#TABLO GELİCEK ALAN
        table=doc.add_table(df1.shape[0]+1, df1.shape[1],style="Light List Accent 1")
        table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.style.font.size=Pt(9)
        if len(desc_list)==1 and len(pn_list)==1 and len(sn_list)==1:
                
                for k in range(df1.shape[-1]):
                        table.cell(0,k).text=df1.columns[k]
                for l in range(df1.shape[0]):
                        for k in range(df1.shape[-1]):
                                table.cell(l+1,k).text=str(df1.values[l,k])
                                
        elif len(desc_list)==2 and len(pn_list)==2 and len(sn_list)==2:
                if desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==3 and len(pn_list)==3 and len(sn_list)==3:
                if desc_list[0]==desc_list[2] and pn_list[0]==pn_list[2] and sn_list[0]==sn_list[2]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif desc_list[0]==desc_list[1] and pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
                
                
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
                
        elif len(desc_list)==4 and len(pn_list)==4 and len(sn_list)==4:
                if desc_list[0]==desc_list[3] and pn_list[0]==pn_list[3] and sn_list[0]==sn_list[3]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[4]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])

                elif  pn_list[0]==pn_list[1] and sn_list[0]==sn_list[1] and pn_list[2]==pn_list[3] and sn_list[2]==sn_list[3]:
                        rowa=table.rows[1]
                        rowb=table.rows[2]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[3]
                        rowd=table.rows[4]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])

                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                        
        elif len(desc_list)==5 and len(pn_list)==5 and len(sn_list)==5:
                if desc_list[0]==desc_list[4] and pn_list[3]==pn_list[4] and sn_list[0]==sn_list[4]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[5]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                elif sn_list[0]==sn_list[2] and sn_list[3]==sn_list[4]:
                        
                        rowa=table.rows[1]
                        rowb=table.rows[3]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        rowc=table.rows[4]
                        rowd=table.rows[5]
                        rowc.cells[0].merge(rowd.cells[0])
                        rowc.cells[1].merge(rowd.cells[1])
                        rowc.cells[2].merge(rowd.cells[2])
                        
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                                        
        elif len(desc_list)==6 and len(pn_list)==6 and len(sn_list)==6:
                if desc_list[0]==desc_list[5] and pn_list[0]==pn_list[5] and sn_list[0]==sn_list[5]:
                        
                
                        rowa=table.rows[1]
                        rowb=table.rows[6]
                        rowa.cells[0].merge(rowb.cells[0])
                        rowa.cells[1].merge(rowb.cells[1])
                        rowa.cells[2].merge(rowb.cells[2])
                        #rowa.cells[3].merge(rowb.cells[3])
                        #rowa.cells[4].merge(rowb.cells[4])
                        #rowa.cells[5].merge(rowb.cells[5])
                
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])
                
        else:
                        for k in range(df1.shape[-1]):
                                table.cell(0,k).text=df1.columns[k]
                        for l in range(df1.shape[0]):
                                for k in range(df1.shape[-1]):
                                        table.cell(l+1,k).text=str(df1.values[l,k])


        for _ in range(2):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10

        chp_style=doc.styles["Body Text"]
        chp=doc.add_paragraph(style=chp_style).add_run(f"{sign}")
        chp.font.size = Pt(11)
        chp.font.name="Arial"
        chp.font.bold= True

        for _ in range(3):
                linespace_style = doc.styles['Body Text']
                linespace = doc.add_paragraph(style=linespace_style).add_run()
                linespace.font.size = 10
        tai_text=f"AJet Hava Taşımacılığı Anonim Şirketi\n        Yeşilköy Mah. Havaalanı Cad. THY Genel Müdürlük No:3/1\n        Bakırköy / İstanbul "
        tai_style=doc.styles["Body Text"]
        tai=doc.add_paragraph(style=tai_style).add_run(f"{tai_text}")
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tai.font.size=Pt(8)
        tai.font.name="Arial"
        tai.font.bold= True
        

        doc.save("word.docx")        
    else:
        print("Error")
    


frame_ust=Frame(master, bg="#EFEFEF")
frame_ust.place(relx=0.1, rely=0.1, relwidth=0.734, relheight=0.1)

frame_alt_sol=Frame(master, bg="#EFEFEF")
frame_alt_sol.place(relx=0.1, rely=0.21, relwidth=0.14, relheight=0.7)
frame_alt_sag=Frame(master,bg="#EFEFEF")
frame_alt_sag.place(relx=0.25, rely=0.21 , relwidth=0.585, relheight=0.7)
nis_tipi_etiket=Label(frame_ust,bg="#EFEFEF", text="NIS TYPE", font="verdana 22 bold")
nis_tipi_etiket.pack(padx=10, side=LEFT)

nis_tipi_opsiyon=tk.StringVar(frame_ust)
nis_tipi_opsiyon.set("\t")
nis_operator_opsiyon=tk.StringVar(frame_ust)
nis_operator_opsiyon.set("\t")

var1="PN/SN/DESC"
var2="BACK TO BIRTH STATEMENT"
var3="INCIDENT/ACCIDENT STATEMENT"
var4="PN/SN/DESC-TSN-CSN"
var5="TURKISH AIRLINES"
var6="AJET"

nis_tipi_acilir_menu=ttk.OptionMenu(
    frame_ust,
    nis_tipi_opsiyon,"Choose A Type",
    var1,
    var2,
    var3,
    var4)
nis_tipi_acilir_menu.pack(padx=10,pady=10,side=LEFT)
nis_operator_menu=ttk.OptionMenu(
        frame_ust,
        nis_operator_opsiyon,"Choose An Operator",
        var5,
        var6)
nis_operator_menu.pack(padx=10,pady=10,side=LEFT)

 

nis_hazirlanis_secici=DateEntry(frame_ust, width=12,background="orange", foreground="black" , borderwidth=1, Locale="tr_TR",date_pattern="dd.mm.yyyy")
nis_hazirlanis_secici._top_cal.overrideredirect(False)
nis_hazirlanis_secici.pack(padx=20, side=RIGHT)

nis_hazırlanıs_tarihi=Label(frame_ust,bg="#EFEFEF", text="CREATE DATE", font="verdana 22 bold")
nis_hazırlanıs_tarihi.pack(padx=20, side=RIGHT)

Label(frame_alt_sol, text="Component Details",bg="#EFEFEF", font="verdana 13 bold").pack(padx=10, pady=10, anchor=N)
Label(frame_alt_sol, text="PN", bg="#EFEFEF", font="verdana 11 bold").pack(padx=10, pady=10,anchor=N)
pn_alani=ttk.Entry(frame_alt_sol, width=20)
pn_alani.pack(padx=10, pady=10, anchor=N)


submit_butonu=ttk.Button(frame_alt_sol, text="Submit", command=getEntry)
submit_butonu.pack(anchor=S)


Label(frame_alt_sol, text="SN", bg="#EFEFEF", font="verdana 11 bold").pack(padx=10, pady=10,anchor=N)
sn_alani=ttk.Entry(frame_alt_sol, width=20)
sn_alani.pack(padx=10, pady=10, anchor=N)


global des_alani


des=Label(frame_alt_sol, text="DESCRIPTION", bg="#EFEFEF", font="verdana 11 bold").pack(padx=10, pady=10,anchor=N)

des_alani=ttk.Entry(frame_alt_sol,width=20)
des_alani.pack(padx=10, pady=10, anchor=N)
des_alani.bind("<Return>", getTransfer)


Label(frame_alt_sol, text="CHAPTER", bg="#EFEFEF", font="verdana 11 bold").pack(padx=10, pady=10,anchor=N)
chp_alani=ttk.Entry(frame_alt_sol, width=20)
chp_alani.pack(padx=10, pady=10, anchor=N)


submit_butonu2=ttk.Button(frame_alt_sol,text="Transfer", command=getTransfer)
submit_butonu2.pack(anchor=S)



link=Label(frame_alt_sol, text="Click for Occurrence Report",font=("Helveticabold",10),fg="#ff1944",cursor="hand2")
link.pack(padx=60, pady=60,anchor=S)
link.bind("<Button-1>", lambda e:
callback("https://occs.thyteknik.com"))





a =2
rows=[]
global trans
trans=["INSTALLED","REMOVED","A/C DELIVERY","A/C RE-DELIVERY"]
def add_row():
        
                                
            global a
            global b
            global items
            global z
            global y
            global x
            global q
            global w
            global tod
            global fd
            global desc
            global pn_area
            global sn_area
            global ac_msn
            global z_top
            global x_bottom
            global q_bottom
            global items  
            global add_button 
            global label_pop
        
                
            a = a + 1
            
            
            items=[]
        
            var = IntVar()
            c = ttk.Checkbutton(frame_alt_sag,variable = var,onvalue=1,offvalue=0)
            c.val=var
            items.append(c)
            c.grid(row = a, column = 0)
            add_button=Button(frame_alt_sag,text="Add",width=4,command=lambda: add(add_button))
            items.append(add_button)
            add_button.grid(row = a, column =0, sticky=E)
            if nis_tipi_opsiyon.get()==var2 or nis_tipi_opsiyon.get()==var3:
                    
                        for j in range(1,10):#Columns

                        
                                b = Entry(frame_alt_sag,width=17,justify="center")
                                items.append(b)
                                b.grid(row=a,column=j)
                                desc=Entry(frame_alt_sag,width=17,justify="center") 
                                items.append(desc)
                                desc.grid(row=a,column=1)
                                pn_area=Entry(frame_alt_sag,width=17,justify="center")
                                items.append(pn_area)
                                pn_area.grid(row=a,column=2)
                                sn_area=Entry(frame_alt_sag,width=17,justify="center")
                                items.append(sn_area)
                                sn_area.grid(row=a,column=3)
                                ac_msn=ttk.Combobox(frame_alt_sag, values=aircraft,width=17,justify="center")
                                ac_msn.grid(row=a,column=4)
                                items.append(ac_msn)
                                z=ttk.Combobox(frame_alt_sag, values=trans,width=17,justify="center")
                                items.append(z)
                                z.grid(row=a-1,column=5)
                                z_top=ttk.Combobox(frame_alt_sag, values=trans,width=17,justify="center")
                                z_top.grid(row=a,column=5)
                                items.append(z_top)
                                tod=Entry(frame_alt_sag,width=17,justify="center")
                                items.append(tod)
                                tod.grid(row=a-1,column=6)
                                fd=Entry(frame_alt_sag,width=17,justify="center")
                                fd.grid(row=a,column=6)
                                items.append(fd)
                                w=Entry(frame_alt_sag,width=17,justify="center")
                                items.append(w)
                                x_bottom=Entry(frame_alt_sag,width=17,justify="center")
                                items.append(x_bottom)
                                x_bottom.grid(row=a,column=7)
                                x=Entry(frame_alt_sag,width=17,justify="center")
                                items.append(x)
                                x.grid(row=a-1,column=7)
                                q_bottom=Entry(frame_alt_sag,width=17,justify="center")
                                items.append(q_bottom)
                                q_bottom.grid(row=a,column=8)
                                q=Entry(frame_alt_sag,width=17,justify="center")
                                items.append(q)
                                q.grid(row=a-1,column=8)






                        rows.append(items)



                        a = a + 1
                        v0 = StringVar()
                        e0 = Entry(frame_alt_sag, textvariable = v0, state = 'readonly',justify='center',width=20)
                        v0.set('Select')
                        e0.grid(row = 1, column = 0)

                        v1 = StringVar()
                        e1 = Entry(frame_alt_sag, textvariable = v1,state = 'readonly',justify='center',width=20)
                        v1.set('DESCRIPTION')
                        e1.grid(row = 1, column = 1 )


                        v2 = StringVar()
                        v2.set('PN')
                        e2 = Entry(frame_alt_sag, textvariable =v2,state="readonly",justify='center',width=20)
                        e2.grid(row = 1, column = 2)


                        v3 = StringVar()
                        e3 = Entry(frame_alt_sag, textvariable = v3,state = 'readonly',justify='center',width=20)
                        v3.set('SN')
                        e3.grid(row = 1, column = 3 )

                        v4 = StringVar()
                        e4 = Entry(frame_alt_sag, textvariable = v4, state = 'readonly',justify='center',width=20)
                        v4.set('A/C_MSN')
                        e4.grid(row = 1, column = 4 )

                        v5 = StringVar()
                        e5 = Entry(frame_alt_sag, textvariable = v5,state = "readonly",justify='center',width=20)
                        v5.set('TRANSACTION')
                        e5.grid(row = 1, column = 5)


                        v6 = StringVar()
                        e6 = Entry(frame_alt_sag, textvariable = v6, state = 'readonly',justify='center',width=20)
                        v6.set('DATE')
                        e6.grid(row = 1, column = 6 )

                        v7 = StringVar()
                        e7 = Entry(frame_alt_sag, textvariable = v7, state = 'readonly',justify='center',width=20)
                        v7.set('TSN')
                        e7.grid(row = 1, column = 7 )

                        v8 = StringVar()
                        e8 = Entry(frame_alt_sag, textvariable = v8, state = 'readonly',justify='center',width=20)
                        v8.set('CSN')
                        e8.grid(row = 1, column = 8 )

            elif nis_tipi_opsiyon.get()==var1:
                    for j in range(1,4):
                            b = Entry(frame_alt_sag,width=17,justify="center")
                            items.append(b)
                            b.grid(row=a,column=j)
                            desc=Entry(frame_alt_sag,width=17,justify="center") 
                            items.append(desc)
                            desc.grid(row=a,column=1)
                            pn_area=Entry(frame_alt_sag,width=17,justify="center")
                            items.append(pn_area)
                            pn_area.grid(row=a,column=2)
                            sn_area=Entry(frame_alt_sag,width=17,justify="center")
                            items.append(sn_area)
                            sn_area.grid(row=a,column=3)

                            tod=Entry(frame_alt_sag,width=17,justify="center")
                            items.append(tod)
                            tod.grid(row=a-1,column=6)
                            fd=Entry(frame_alt_sag,width=17,justify="center")
                            fd.grid(row=a,column=6)
                            items.append(fd)

                            w=Entry(frame_alt_sag,width=17,justify="center")
                            items.append(w)
                            
                    rows.append(items)
                    a=a+1
                    v0 = StringVar()
                    e0 = Entry(frame_alt_sag, textvariable = v0, state = 'readonly',justify='center',width=20)
                    v0.set('Select')
                    e0.grid(row = 1, column = 0)    
                    v1 = StringVar()
                    e1 = Entry(frame_alt_sag, textvariable = v1,state = 'readonly',justify='center',width=20)
                    v1.set('DESCRIPTION')
                    e1.grid(row = 1, column = 1 )    
                    v2 = StringVar()
                    v2.set('PN')
                    e2 = Entry(frame_alt_sag, textvariable =v2,state="readonly",justify='center',width=20)
                    e2.grid(row = 1, column = 2)    
                    v3 = StringVar()
                    e3 = Entry(frame_alt_sag, textvariable = v3,state = 'readonly',justify='center',width=20)
                    v3.set('SN')
                    e3.grid(row = 1, column = 3 )
                    

                    v6 = StringVar()
                    e6 = Entry(frame_alt_sag, textvariable = v6, state = 'readonly',justify='center',width=20)
                    v6.set('DATE')
                    e6.grid(row = 1, column = 6 )
                            
                            
                    
            elif nis_tipi_opsiyon.get()==var4:
                    
                    for j in range(1,4):
                            b = Entry(frame_alt_sag,width=17,justify="center")
                            items.append(b)
                            b.grid(row=a,column=j)
                            desc=Entry(frame_alt_sag,width=17,justify="center") 
                            items.append(desc)
                            desc.grid(row=a,column=1)
                            pn_area=Entry(frame_alt_sag,width=17,justify="center")
                            items.append(pn_area)
                            pn_area.grid(row=a,column=2)
                            sn_area=Entry(frame_alt_sag,width=17,justify="center")
                            items.append(sn_area)
                            sn_area.grid(row=a,column=3)
                            tod=Entry(frame_alt_sag,width=17,justify="center")
                            items.append(tod)
                            tod.grid(row=a-1,column=6)
                            fd=Entry(frame_alt_sag,width=17,justify="center")
                            fd.grid(row=a,column=6)
                            items.append(fd)
                            w=Entry(frame_alt_sag,width=17,justify="center")
                            items.append(w)
                            x_bottom=Entry(frame_alt_sag,width=17,justify="center")
                            items.append(x_bottom)
                            x_bottom=Entry(frame_alt_sag,width=17,justify="center")
                            items.append(x_bottom)
                            x_bottom.grid(row=a,column=7)
                            q_bottom=Entry(frame_alt_sag,width=17,justify="center")
                            items.append(q_bottom)
                            q_bottom.grid(row=a,column=8)
                    rows.append(items)
                    a=a+1
                    v0 = StringVar()
                    e0 = Entry(frame_alt_sag, textvariable = v0, state = 'readonly',justify='center',width=20)
                    v0.set('Select')
                    e0.grid(row = 1, column = 0)    
                    v1 = StringVar()
                    e1 = Entry(frame_alt_sag, textvariable = v1,state = 'readonly',justify='center',width=20)
                    v1.set('DESCRIPTION')
                    e1.grid(row = 1, column = 1 )    
                    v2 = StringVar()
                    v2.set('PN')
                    e2 = Entry(frame_alt_sag, textvariable =v2,state="readonly",justify='center',width=20)
                    e2.grid(row = 1, column = 2)    
                    v3 = StringVar()
                    e3 = Entry(frame_alt_sag, textvariable = v3,state = 'readonly',justify='center',width=20)
                    v3.set('SN')
                    e3.grid(row = 1, column = 3 )
                    v6 = StringVar()
                    e6 = Entry(frame_alt_sag, textvariable = v6, state = 'readonly',justify='center',width=20)
                    v6.set('DATE')
                    e6.grid(row = 1, column = 6 )
                    v7 = StringVar()
                    e7 = Entry(frame_alt_sag, textvariable = v7, state = 'readonly',justify='center',width=20)
                    v7.set('TSN')
                    e7.grid(row = 1, column = 7 )    
                    v8 = StringVar()
                    e8 = Entry(frame_alt_sag, textvariable = v8, state = 'readonly',justify='center',width=20)
                    v8.set('CSN')
                    e8.grid(row = 1, column = 8 )                          
#labels=[]                            
def delete_row():
        for rowno, row in reversed(list(enumerate(rows))):
                if row[0].val.get()==1:
                        pn_list.clear()
                        desc_list.clear()
                        sn_list.clear()
                        ac_list.clear()
                        z_top_list.clear()
                        fd_list.clear()
                        tsn_list.clear()
                        csn_list.clear()
                        #labels.clear()
                        for a in row:
                                a.destroy() 
                        rows.pop(rowno)
                        #for label_pop in labels:  # Assuming you have a list of labels
                        #        if label_pop.cget('text') == "INCIDENT/ACCIDENT":
                        #                label_pop.destroy()  # Destroy the label
                        #                labels.remove(label_pop)
bt = ttk.Button(frame_alt_sag , text = 'Add Row', command = add_row)
bt.grid(row =0, column=0)


dl =ttk.Button(frame_alt_sag , text = 'Delete Row',command= delete_row)
dl.grid(row =0, column=1)
global e1
global v1
c1=ttk.Button(frame_alt_sag , text ="Convert To Doc",command=convert)
c1.grid(row =0,column=8)
#İmza için tuş.
method = tk.StringVar(frame_alt_sag)
method.set("\t")
s1="Furkan ULCAY"
s2="Bayram Erkan ÇANKAYA"
s3="Özkan AYDOĞDU"
s4="AJET"
sign_menu=ttk.OptionMenu(frame_alt_sag,method,"Signed By",s1,s2,s3,s4)
sign_menu.grid(row=0,column=6)



master.mainloop()


